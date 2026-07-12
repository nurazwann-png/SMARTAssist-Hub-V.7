"""Report Generator Agent — Section 4: Bureaucratic Reporting Specialist."""

import json
import os
import re
import io
import pathlib
from datetime import datetime
from backend.deepseek_client import chat_completion

# ── Field schema ──

REPORT_FIELDS = [
    {"key": "nama_program", "label": "Nama Program", "example": "Bengkel ICT Guru Besar Siri 2/2026"},
    {"key": "tarikh_program", "label": "Tarikh Program", "example": "10 Julai 2026"},
    {"key": "hari", "label": "Hari", "example": "Khamis"},
    {"key": "masa", "label": "Masa Program", "example": "8.00 pagi - 5.00 petang"},
    {"key": "organisasi", "label": "Nama Organisasi", "example": "Pejabat Pendidikan Daerah Dalat"},
    {"key": "pegawai_nama", "label": "Nama Pegawai Yang Terlibat", "example": "Ahmad bin Ali, Siti binti Hassan, Razak bin Omar"},
    {"key": "pegawai_jawatan", "label": "Jawatan Pegawai Yang Terlibat", "example": "Pegawai Pendidikan Daerah, Penolong Pegawai Pendidikan, Guru Besar"},
    {"key": "objektif", "label": "Objektif Program", "example": "Meningkatkan kemahiran ICT guru besar"},
    {"key": "rumusan", "label": "Rumusan / Laporan Ringkas", "example": "Program berjalan lancar dengan 45 peserta"},
    {"key": "cadangan", "label": "Cadangan / Tindakan Susulan", "example": "Mengadakan bengkel susulan pada bulan hadapan"},
    {"key": "penyedia_nama", "label": "Nama Penyedia Laporan", "example": "Mohd Razali bin Ibrahim"},
    {"key": "penyedia_jawatan", "label": "Jawatan Penyedia", "example": "Pegawai Pendidikan Daerah"},
    {"key": "tarikh_disediakan", "label": "Tarikh Disediakan", "example": "11 Julai 2026"},
    {"key": "pengesah_nama", "label": "Nama Pengesah Laporan", "example": "Dato' Ahmad bin Yusof"},
    {"key": "pengesah_jawatan", "label": "Jawatan Pengesah", "example": "Pegawai Pendidikan Daerah Kanan"},
]

AUTO_GENERATE_KEYS = {"rumusan", "cadangan"}

_SYSTEM_PROMPT_TEMPLATE = """Anda ialah pakar laporan birokrasi dalam SMARTAssist Hub.
Tugas anda: membantu pengguna menyediakan laporan satu muka surat (One Page Report) dalam format rasmi PPD/KPM.

CARA KERJA:
1. Tanya maklumat yang diperlukan SATU DEMI SATU — HANYA SATU SOALAN setiap giliran
2. Jika pengguna beri beberapa maklumat sekaligus, terima semua yang diberi, kemudian tanya field seterusnya yang masih kosong
3. JANGAN tanya "rumusan" dan "cadangan" — anda WAJIB menjana kedua-duanya secara automatik berdasarkan maklumat program yang dikumpul
4. Jika maklumat tidak mencukupi untuk menjana rumusan/cadangan yang bermakna, tanya soalan spesifik (cth: "Berapa ramai peserta yang hadir?" atau "Apakah aktiviti utama yang dijalankan?")

URUTAN SOALAN:
nama_program → tarikh_program → hari → masa → organisasi → pegawai_nama → pegawai_jawatan → objektif → (jana rumusan & cadangan automatik) → penyedia_nama → penyedia_jawatan → tarikh_disediakan → pengesah_nama → pengesah_jawatan

GAYA:
- Formal, tiada hiasan, seperti pegawai kerajaan yang berpengalaman
- Gunakan Bahasa Malaysia rasmi
- HANYA tanya SATU field pada satu masa
- Rumusan dan cadangan dijana dalam gaya laporan rasmi — perenggan bernombor, bahasa formal

TARIKH SEMASA: {current_date}
Tahun semasa ialah {current_year}. Jangan persoalkan tarikh yang pengguna berikan.

FORMAT OUTPUT — balas HANYA dalam JSON:
{{
  "phase": 0-3,
  "message": "Mesej kepada pengguna",
  "field_asking": "nama field yang sedang ditanya (atau null)",
  "fields_collected": {{"key": "value", ...}},
  "document_preview": "Teks laporan lengkap (hanya pada phase 2-3)",
  "validation_errors": ["Senarai ralat jika ada"],
  "ready_to_save": true/false
}}

PHASES:
- Phase 0: Pengenalan — sahkan pengguna mahu buat One Page Report
- Phase 1: Kumpul maklumat secara berperingkat. JANGAN tanya rumusan dan cadangan.
- Phase 2: Jana rumusan dan cadangan secara automatik, tunjukkan pratonton laporan lengkap. Selepas itu tanya penyedia_nama, penyedia_jawatan, tarikh_disediakan, pengesah_nama, pengesah_jawatan jika belum ada.
- Phase 3: Laporan disahkan dan sedia untuk dimuat turun

FIELD KEYS WAJIB (guna key tepat ini dalam fields_collected):
nama_program, tarikh_program, hari, masa, organisasi, pegawai_nama, pegawai_jawatan, objektif, rumusan, cadangan, penyedia_nama, penyedia_jawatan, tarikh_disediakan, pengesah_nama, pengesah_jawatan

PENTING:
- Jangan reka maklumat peribadi — tanya pengguna
- RUMUSAN dan CADANGAN WAJIB dijana oleh anda secara automatik berdasarkan nama program, objektif, dan maklumat lain
- Gunakan KEY TEPAT seperti senarai di atas dalam fields_collected
- Pastikan JSON sah"""

MAX_IMAGES = 4
_IMAGES_DIR = pathlib.Path("static/report_images")
_NS = "report"  # namespace dalam SessionStore
_NS_IMG = "report_images"


def _validate_landscape(image_bytes: bytes, filename: str) -> str | None:
    try:
        from PIL import Image as PILImage
        with PILImage.open(io.BytesIO(image_bytes)) as img:
            w, h = img.size
            if w <= h:
                return f"Gambar '{filename}' ({w}×{h}px) mesti landscape (lebar > tinggi)."
    except Exception:
        pass
    return None


def add_report_image(session_id: str, image_bytes: bytes, filename: str) -> dict:
    from backend.session_store import get_store
    _IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    err = _validate_landscape(image_bytes, filename)
    if err:
        return {"ok": False, "error": err}
    store = get_store()
    images = store.get(session_id, _NS_IMG, "list", [])
    if len(images) >= MAX_IMAGES:
        return {"ok": False, "error": f"Maksimum {MAX_IMAGES} gambar sahaja dibenarkan."}
    safe_name = re.sub(r'[^\w\-_\.]', '_', f"{session_id}_{len(images)}_{filename}")
    file_path = _IMAGES_DIR / safe_name
    file_path.write_bytes(image_bytes)
    images.append({"filename": filename, "path": str(file_path), "safe_name": safe_name})
    store.set(session_id, _NS_IMG, "list", images)
    return {"ok": True, "index": len(images) - 1, "count": len(images), "max": MAX_IMAGES}


def remove_report_image(session_id: str, index: int) -> dict:
    from backend.session_store import get_store
    store = get_store()
    images = store.get(session_id, _NS_IMG, "list", [])
    if 0 <= index < len(images):
        img = images.pop(index)
        try:
            pathlib.Path(img["path"]).unlink(missing_ok=True)
        except Exception:
            pass
        store.set(session_id, _NS_IMG, "list", images)
        return {"ok": True, "count": len(images)}
    return {"ok": False, "error": "Indeks tidak sah"}


def get_report_images(session_id: str) -> list[dict]:
    from backend.session_store import get_store
    return get_store().get(session_id, _NS_IMG, "list", [])


def clear_report_images(session_id: str):
    from backend.session_store import get_store
    images = get_store().get(session_id, _NS_IMG, "list", [])
    for img in images:
        try:
            pathlib.Path(img["path"]).unlink(missing_ok=True)
        except Exception:
            pass
    get_store().delete_ns(session_id, _NS_IMG)


def _get_session(session_id: str) -> dict:
    from backend.session_store import get_store
    data = get_store().get_all(session_id, _NS)
    if not data or "phase" not in data:
        default = {"phase": 0, "fields": {}, "document": None}
        get_store().set_all(session_id, _NS, default)
        return default
    return data


def _save_session(session_id: str, session: dict):
    from backend.session_store import get_store
    get_store().set_all(session_id, _NS, session)


def _find_missing_fields(collected: dict) -> list[dict]:
    return [f for f in REPORT_FIELDS if f["key"] not in collected or not collected[f["key"]]]


def _auto_generate_rumusan_cadangan(fields: dict) -> tuple[str, str]:
    """Jana rumusan dan cadangan secara automatik berdasarkan maklumat program."""
    prompt = f"""Anda adalah pegawai kerajaan. Tulis rumusan pelaksanaan dan cadangan tindakan susulan dalam Bahasa Malaysia rasmi berdasarkan maklumat berikut.

Nama Program: {fields.get('nama_program', '')}
Tarikh: {fields.get('tarikh_program', '')}
Organisasi: {fields.get('organisasi', '')}
Pegawai Terlibat: {fields.get('pegawai_nama', '')}
Objektif: {fields.get('objektif', '')}

Balas dalam format JSON tepat ini sahaja:
{{"rumusan": "...", "cadangan": "..."}}

Rumusan: 2-3 ayat padat tentang pelaksanaan program.
Cadangan: 2-3 cadangan tindakan susulan yang konkrit.
Bahasa Malaysia rasmi. Tiada placeholder."""
    try:
        raw = chat_completion(messages=[
            {"role": "system", "content": "Balas dalam format JSON sahaja. Tiada markdown."},
            {"role": "user", "content": prompt}
        ], temperature=0.3, max_tokens=500)
        data = _try_parse_json(raw)
        if data and data.get("rumusan"):
            return data.get("rumusan", ""), data.get("cadangan", "")
    except Exception:
        pass
    # Fallback ringkas jika LLM gagal
    nama = fields.get('nama_program', 'program')
    org = fields.get('organisasi', 'organisasi')
    obj = fields.get('objektif', '')
    rumusan = f"Program {nama} telah dilaksanakan dengan jayanya oleh {org}. Semua objektif program telah dicapai dengan penglibatan penuh daripada semua peserta. {obj}."
    cadangan = f"Program susulan perlu diadakan bagi memastikan kesinambungan pencapaian objektif. Pemantauan berterusan perlu dilakukan oleh {org}. Laporan program ini perlu dikemukakan kepada pihak atasan untuk tindakan selanjutnya."
    return rumusan, cadangan


def _has_placeholders(text: str) -> list[str]:
    patterns = [r'\[PLACEHOLDER\]', r'\{\{CREATIVE:[^}]+\}\}', r'\{\{DYNAMIC_LIST:[^}]+\}\}']
    found = []
    for p in patterns:
        found.extend(re.findall(p, text))
    return found


def _build_report(f: dict) -> str:
    def _split_field(val):
        if isinstance(val, list):
            return [x.strip() for x in val if str(x).strip()]
        return [x.strip() for x in str(val).split(",") if x.strip()]

    nama_list = _split_field(f.get("pegawai_nama", ""))
    jawatan_list = _split_field(f.get("pegawai_jawatan", ""))
    if nama_list:
        rows = []
        for i, nama in enumerate(nama_list):
            jawatan = jawatan_list[i] if i < len(jawatan_list) else ""
            rows.append(f"{i+1}) {nama}" + (f"\n   {jawatan}" if jawatan else ""))
        pegawai_lines = "\n".join(rows)
    else:
        pegawai_lines = "[PLACEHOLDER]"

    objektif = f.get("objektif", "[PLACEHOLDER]")
    if isinstance(objektif, list):
        objektif_lines = "\n".join(f"{i+1}) {o.strip()}" for i, o in enumerate(objektif))
    else:
        items = [o.strip() for o in objektif.split(",") if o.strip()]
        objektif_lines = "\n".join(f"{i+1}) {o}" for i, o in enumerate(items)) if items else objektif

    return f"""ONE PAGE REPORT

NAMA PROGRAM
{f.get('nama_program', '[PLACEHOLDER]')}

BUTIRAN PERLAKSANAAN
Tarikh      : {f.get('tarikh_program', '[PLACEHOLDER]')}
Hari        : {f.get('hari', '[PLACEHOLDER]')}
Masa        : {f.get('masa', '[PLACEHOLDER]')}
Organisasi  : {f.get('organisasi', '[PLACEHOLDER]')}

PEGAWAI YANG TERLIBAT
{pegawai_lines}

OBJEKTIF
{objektif_lines}

RUMUSAN / LAPORAN
{f.get('rumusan', '[PLACEHOLDER]')}

CADANGAN / TINDAKAN
{f.get('cadangan', '[PLACEHOLDER]')}

DISEDIAKAN OLEH                          DISAHKAN OLEH

............................................................    ............................................................
NAMA    : {f.get('penyedia_nama', '[PLACEHOLDER]')}              NAMA    : {f.get('pengesah_nama', '[PLACEHOLDER]')}
JAWATAN : {f.get('penyedia_jawatan', '[PLACEHOLDER]')}           JAWATAN : {f.get('pengesah_jawatan', '[PLACEHOLDER]')}
TARIKH  : {f.get('tarikh_disediakan', '[PLACEHOLDER]')}"""


def handle(query: str, history: list[dict] | None = None, session_id: str = "default", lang: str = "bm") -> str:
    if query == '__INTRO__':
        if lang == "en":
            return ("Assalamualaikum and welcome! 📋 I am the Report Generator. I am ready to help you generate a one-page report (One Page Report) in the official PPD/KPM format. Could you tell me the name of the programme or activity to be reported?\n\n"
                    "⚠️ Reminder: All reports generated by this system are AI-produced drafts. Please review all content, statistics and conclusions carefully before official submission.")
        return ("Assalamualaikum dan selamat datang! 📋 Saya Penjana Laporan. Saya sedia membantu tuan/puan menjana laporan satu muka surat (One Page Report) dalam format rasmi PPD/KPM. Boleh beritahu saya nama program atau aktiviti yang ingin dilaporkan?\n\n"
                "⚠️ Peringatan: Semua laporan yang dijana adalah draf hasil AI. Sila semak semua kandungan, statistik dan rumusan dengan teliti sebelum dikemukakan secara rasmi.")

    session = _get_session(session_id)

    _PATCH_KEYWORDS = {"kemaskini", "ubah", "ganti", "tukar", "edit", "update", "betulkan", "perbetulkan", "ubah suai"}
    query_lower = query.lower()
    if session.get("document") and any(kw in query_lower for kw in _PATCH_KEYWORDS):
        patch_prompt = f"""Laporan semasa:
---
{session['document']}
---

Pengguna meminta perubahan berikut: {query}

Kembalikan HANYA laporan yang telah dikemaskini dalam format JSON:
{{
  "phase": {session['phase']},
  "message": "Laporan telah dikemaskini.",
  "document_preview": "<laporan lengkap selepas perubahan>",
  "ready_to_save": true
}}"""
        now = datetime.now()
        date_str = now.strftime("%#d %B %Y") if os.name == "nt" else now.strftime("%-d %B %Y")
        system_prompt = _SYSTEM_PROMPT_TEMPLATE.replace("{current_date}", date_str).replace("{current_year}", str(now.year))
        _patch_lang_note = "\n\nIMPORTANT: The user has selected English. Respond in English for all 'message' fields." if lang == "en" else ""
        patch_messages = [{"role": "system", "content": system_prompt + _patch_lang_note}, {"role": "user", "content": patch_prompt}]
        try:
            raw = chat_completion(messages=patch_messages, temperature=0.2, max_tokens=3000)
            parsed = _try_parse_json(raw)
            if parsed and parsed.get("document_preview"):
                doc_text = parsed["document_preview"]
                session["document"] = doc_text
                parsed["ready_to_save"] = not bool(_has_placeholders(doc_text))
                image_count = len(get_report_images(session_id))
                parsed["awaiting_images"] = True
                parsed["image_count"] = image_count
                parsed["max_images"] = MAX_IMAGES
                return json.dumps(parsed, ensure_ascii=False)
        except Exception:
            pass

    missing = _find_missing_fields(session["fields"])
    context_info = f"""
Status sesi semasa:
- Phase: {session['phase']}
- Field terkumpul: {json.dumps(session['fields'], ensure_ascii=False) if session['fields'] else 'tiada lagi'}
- Field belum diisi: {json.dumps([f['label'] for f in missing], ensure_ascii=False)}
"""

    now = datetime.now()
    date_str = now.strftime("%#d %B %Y") if os.name == "nt" else now.strftime("%-d %B %Y")
    system_prompt = _SYSTEM_PROMPT_TEMPLATE.replace("{current_date}", date_str).replace("{current_year}", str(now.year))
    lang_note = "\n\nIMPORTANT: The user has selected English. You MUST respond entirely in English. All 'message' and conversational text fields in your JSON must be in English. The generated report document content should remain in Malay as it is an official KPM document." if lang == "en" else ""

    messages = [
        {"role": "system", "content": system_prompt + context_info + lang_note},
    ]
    if history:
        for msg in history[-10:]:
            messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": query})

    try:
        raw = chat_completion(messages=messages, temperature=0.3, max_tokens=3000)
    except RuntimeError as e:
        return json.dumps({
            "phase": session["phase"],
            "message": f"Ralat menghubungi API. Sila cuba lagi.\n{e}",
            "ready_to_save": False,
        }, ensure_ascii=False)

    parsed = _try_parse_json(raw)
    if not parsed:
        return json.dumps({
            "phase": session["phase"],
            "message": raw,
            "ready_to_save": False,
        }, ensure_ascii=False)

    if parsed.get("fields_collected"):
        session["fields"].update(parsed["fields_collected"])

    if parsed.get("phase") is not None:
        session["phase"] = parsed["phase"]

    # Auto-jana rumusan/cadangan apabila field program utama sudah lengkap
    _PROGRAM_KEYS = {"nama_program", "tarikh_program", "hari", "masa", "organisasi", "pegawai_nama", "pegawai_jawatan", "objektif"}
    f = session["fields"]
    program_complete = all(f.get(k) for k in _PROGRAM_KEYS)
    if program_complete and (not f.get("rumusan") or not f.get("cadangan")):
        rumusan, cadangan = _auto_generate_rumusan_cadangan(f)
        if rumusan:
            f["rumusan"] = rumusan
        if cadangan:
            f["cadangan"] = cadangan

    all_filled = not _find_missing_fields(session["fields"])
    if session["phase"] >= 2 or all_filled:
        if all_filled:
            session["phase"] = max(session["phase"], 2)
            parsed["phase"] = session["phase"]
        doc_text = _build_report(session["fields"])
        placeholders = _has_placeholders(doc_text)
        parsed["document_preview"] = doc_text
        parsed["document_html"] = _build_report_html(session["fields"])
        if placeholders:
            parsed["validation_errors"] = [f"Masih ada placeholder yang belum diisi: {', '.join(placeholders)}"]
            parsed["ready_to_save"] = False
        else:
            parsed["ready_to_save"] = True
            session["document"] = doc_text
            # Prompt for images
            image_count = len(get_report_images(session_id))
            parsed["awaiting_images"] = True
            parsed["image_count"] = image_count
            parsed["max_images"] = MAX_IMAGES
            if image_count == 0:
                parsed["message"] = (
                    parsed.get("message", "Laporan telah siap!") +
                    "\n\n📷 **Langkah Akhir:** Laporan anda telah siap! Sila muat naik sehingga 4 gambar landscape untuk dilampirkan dalam laporan. Gunakan butang '+ Tambah Gambar' di bawah, atau terus muat turun jika tiada gambar diperlukan."
                )

    parsed["fields_status"] = {
        "collected": session["fields"],
        "missing": [f["label"] for f in _find_missing_fields(session["fields"])],
    }

    _save_session(session_id, session)
    return json.dumps(parsed, ensure_ascii=False)


def _build_report_html(f: dict) -> str:
    logo_url = None
    try:
        from backend.letterhead_store import get_active_by_type
        lh = get_active_by_type("logo")
        if lh and lh.get("filename"):
            logo_url = f"/api/letterhead/image/{lh['filename']}"
    except Exception:
        pass

    if logo_url:
        logo_block = f'<img src="{logo_url}" style="max-width:100%;max-height:100px;display:block;margin:0 auto">'
    else:
        logo_block = '<div style="height:60px;background:#f5f5f5;display:flex;align-items:center;justify-content:center;color:#aaa;font-size:10pt;border:1px dashed #ccc">[Logo]</div>'

    def _split_f(val):
        if isinstance(val, list):
            return [x.strip() for x in val if str(x).strip()]
        return [x.strip() for x in str(val).split(",") if x.strip()]

    nama_pg = _split_f(f.get("pegawai_nama", ""))
    jawatan_pg = _split_f(f.get("pegawai_jawatan", ""))
    pg_parts = []
    for i, nm in enumerate(nama_pg):
        jw = jawatan_pg[i] if i < len(jawatan_pg) else ""
        pg_parts.append(f"{i+1}) {nm}" + (f"<br>&nbsp;&nbsp;&nbsp;{jw}" if jw else ""))
    pegawai_html = "<br>".join(pg_parts) if pg_parts else ""

    objektif = f.get("objektif", "")
    if isinstance(objektif, list):
        obj_items = objektif
    else:
        obj_items = [o.strip() for o in objektif.split(",") if o.strip()]
    obj_html = "<br>".join(f"{i+1}) {o}" for i, o in enumerate(obj_items)) if obj_items else objektif

    rumusan = f.get("rumusan", "").replace("\n", "<br>")
    cadangan = f.get("cadangan", "").replace("\n", "<br>")

    TH = 'style="border:1px solid #000;padding:5px 8px;font-weight:bold"'
    TC = 'style="border:1px solid #000;padding:5px 8px;vertical-align:top"'
    TL = 'style="border:1px solid #000;padding:5px 8px;width:20%"'
    TV = 'style="border:1px solid #000;padding:5px 8px;width:30%"'

    return (
        f'<div style="font-family:Arial,sans-serif;font-size:10pt;line-height:1.5;color:#000">'
        f'{logo_block}'
        f'<hr style="border:none;border-top:2px solid #000;margin:8px 0">'
        f'<table style="width:100%;border-collapse:collapse">'
        f'<tr><td colspan="4" style="border:1px solid #000;padding:6px 8px;text-align:center;font-weight:bold;font-size:11pt">ONE PAGE REPORT</td></tr>'
        f'<tr><td colspan="4" {TH}>NAMA PROGRAM</td></tr>'
        f'<tr><td colspan="4" {TC}>{f.get("nama_program","")}</td></tr>'
        f'<tr><td colspan="4" {TH}>BUTIRAN PERLAKSANAAN</td></tr>'
        f'<tr>'
        f'<td {TL}>Tarikh</td><td {TV}>{f.get("tarikh_program","")}</td>'
        f'<td {TL}>Hari</td><td {TV}>{f.get("hari","")}</td>'
        f'</tr>'
        f'<tr>'
        f'<td {TL}>Masa</td><td {TV}>{f.get("masa","")}</td>'
        f'<td {TL}>Nama Sekolah</td><td {TV}>{f.get("organisasi","")}</td>'
        f'</tr>'
        f'<tr><td colspan="4" {TH}>PEGAWAI YANG TERLIBAT</td></tr>'
        f'<tr><td colspan="4" style="border:1px solid #000;padding:5px 8px;min-height:50px;vertical-align:top">{pegawai_html}</td></tr>'
        f'<tr><td colspan="4" {TH}>OBJEKTIF</td></tr>'
        f'<tr><td colspan="4" style="border:1px solid #000;padding:5px 8px;min-height:50px;vertical-align:top">{obj_html}</td></tr>'
        f'<tr><td colspan="4" {TH}>RUMUSAN / LAPORAN</td></tr>'
        f'<tr><td colspan="4" style="border:1px solid #000;padding:5px 8px;min-height:70px;vertical-align:top">{rumusan}</td></tr>'
        f'<tr><td colspan="4" {TH}>CADANGAN / TINDAKAN</td></tr>'
        f'<tr><td colspan="4" style="border:1px solid #000;padding:5px 8px;min-height:70px;vertical-align:top">{cadangan}</td></tr>'
        f'<tr>'
        f'<td colspan="2" {TH}>DISEDIAKAN OLEH</td>'
        f'<td colspan="2" style="border:1px solid #000;padding:5px 8px;font-weight:bold;text-align:center">DISAHKAN OLEH</td>'
        f'</tr>'
        f'<tr>'
        f'<td colspan="2" style="border:1px solid #000;padding:8px;height:80px;vertical-align:bottom">'
        f'............................................................<br>'
        f'NAMA &nbsp;&nbsp;&nbsp;: {f.get("penyedia_nama","")}<br>'
        f'JAWATAN : {f.get("penyedia_jawatan","")}<br>'
        f'TARIKH &nbsp;&nbsp;: {f.get("tarikh_disediakan","")}'
        f'</td>'
        f'<td colspan="2" style="border:1px solid #000;padding:8px;height:80px;vertical-align:bottom">'
        f'............................................................<br>'
        f'NAMA &nbsp;&nbsp;&nbsp;: {f.get("pengesah_nama","")}<br>'
        f'JAWATAN : {f.get("pengesah_jawatan","")}'
        f'</td>'
        f'</tr>'
        f'</table>'
        f'</div>'
    )


def get_document(session_id: str) -> str | None:
    session = _get_session(session_id)
    if session and session.get("document"):
        return session["document"]
    return None


def get_fields(session_id: str) -> dict:
    session = _get_session(session_id)
    return session.get("fields", {}).copy() if session else {}


def apply_improvement(session_id: str, improved_fields: dict) -> str | None:
    """Update session fields and rebuild document using original template. Returns new doc text."""
    session = _get_session(session_id)
    if not session or not session.get("fields"):
        return None
    for key in ("rumusan", "cadangan"):
        if key in improved_fields and improved_fields[key]:
            session.setdefault("fields", {})[key] = improved_fields[key]
    # Rebuild document from updated fields using original template
    new_doc = _build_report(session["fields"])
    session["document"] = new_doc
    _save_session(session_id, session)
    return new_doc


GENERATED_FIELD_KEYS = ["rumusan", "cadangan"]


def get_session_info(session_id: str) -> dict | None:
    s = _get_session(session_id)
    return s if s.get("fields") else None


def build_docx(session_id: str) -> bytes | None:
    session = _get_session(session_id)
    if not session or not session.get("document"):
        return None

    f = session["fields"]

    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    def _run(para, text, bold=False, size=11, center=False):
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.bold = bold
        return r

    def _set_cell_border(cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
            + ''.join(f'<w:{edge} w:val="{v["val"]}" w:sz="{v["sz"]}" w:space="0" w:color="{v["color"]}"/>'
                      for edge, v in kwargs.items())
            + '</w:tcBorders>')
        tcPr.append(tcBorders)

    def _border_kwargs():
        b = {"val": "single", "sz": "4", "color": "000000"}
        return {"top": b, "bottom": b, "start": b, "end": b}

    def _shade_cell(cell, color="FFC000"):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # ── Header: letterhead image or fallback KPM text ──
    _lh_inserted = False
    try:
        from backend.letterhead_store import get_active_path_by_type
        _lh_path = get_active_path_by_type("logo")
        if _lh_path:
            lh_para = doc.add_paragraph()
            lh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lh_para.paragraph_format.space_after = Pt(6)
            lh_para.add_run().add_picture(str(_lh_path), width=Cm(16))
            _lh_inserted = True
    except Exception:
        pass

    if not _lh_inserted:
        h1 = doc.add_paragraph()
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(h1, "KEMENTERIAN PENDIDIKAN MALAYSIA", bold=True, size=12, center=True)
        h2 = doc.add_paragraph()
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(h2, f"{f.get('organisasi', 'Pejabat Pendidikan Daerah')}", bold=False, size=10, center=True)

    # Horizontal line
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(4)
    line_para.paragraph_format.space_after = Pt(8)
    pPr = line_para._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/></w:pBdr>')
    pPr.append(pBdr)

    total_width = Cm(15.92)  # A4 - margins

    # ── Main table ──
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Cm(3.0), Cm(4.96), Cm(3.0), Cm(4.96)]

    def _add_full_row(text, bold=True, shading=None, center=True, min_height=None):
        row = table.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[3])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, text, bold=bold)
        _set_cell_border(cell, **_border_kwargs())
        if shading:
            _shade_cell(cell, shading)
        if min_height:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{min_height}" w:hRule="atLeast"/>')
            trPr.append(trHeight)
        return cell

    def _add_two_col_row(left_text, right_text, left_bold=True, right_bold=False, right_center=False, min_height=None):
        row = table.add_row()
        left_cell = row.cells[0]
        left_cell.merge(row.cells[1])
        right_cell = row.cells[2]
        right_cell.merge(row.cells[3])

        lp = left_cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after = Pt(4)
        _run(lp, left_text, bold=left_bold)
        _set_cell_border(left_cell, **_border_kwargs())

        rp = right_cell.paragraphs[0]
        rp.paragraph_format.space_before = Pt(4)
        rp.paragraph_format.space_after = Pt(4)
        if right_center:
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(rp, right_text, bold=right_bold)
        _set_cell_border(right_cell, **_border_kwargs())

        if min_height:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{min_height}" w:hRule="atLeast"/>')
            trPr.append(trHeight)
        return left_cell, right_cell

    # Row 1: ONE PAGE REPORT
    _add_full_row("ONE PAGE REPORT", bold=True)

    # Row 2: NAMA PROGRAM + value
    _add_full_row("NAMA PROGRAM", bold=True)
    _add_full_row(f.get("nama_program", ""), bold=False)

    # Row 3: BUTIRAN PERLAKSANAAN header
    _add_full_row("BUTIRAN PERLAKSANAAN", bold=True)

    # Row 4: Tarikh | value | Hari | value
    row4 = table.add_row()
    cells4 = row4.cells
    for i, (label, val) in enumerate([
        ("Tarikh", f.get("tarikh_program", "")),
        ("Hari", f.get("hari", "")),
    ]):
        label_cell = cells4[i * 2]
        val_cell = cells4[i * 2 + 1]
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after = Pt(4)
        _run(lp, label, bold=False)
        _set_cell_border(label_cell, **_border_kwargs())

        vp = val_cell.paragraphs[0]
        vp.paragraph_format.space_before = Pt(4)
        vp.paragraph_format.space_after = Pt(4)
        _run(vp, val, bold=False)
        _set_cell_border(val_cell, **_border_kwargs())

    # Row 5: Masa | value | Nama Sekolah | value
    row5 = table.add_row()
    cells5 = row5.cells
    for i, (label, val) in enumerate([
        ("Masa", f.get("masa", "")),
        ("Nama Sekolah", f.get("organisasi", "")),
    ]):
        label_cell = cells5[i * 2]
        val_cell = cells5[i * 2 + 1]
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after = Pt(4)
        _run(lp, label, bold=False)
        _set_cell_border(label_cell, **_border_kwargs())

        vp = val_cell.paragraphs[0]
        vp.paragraph_format.space_before = Pt(4)
        vp.paragraph_format.space_after = Pt(4)
        _run(vp, val, bold=False)
        _set_cell_border(val_cell, **_border_kwargs())

    # Row 6: PEGAWAI YANG TERLIBAT header + numbered list
    _add_full_row("PEGAWAI YANG TERLIBAT", bold=True)
    def _split_f(val):
        if isinstance(val, list):
            return [x.strip() for x in val if str(x).strip()]
        return [x.strip() for x in str(val).split(",") if x.strip()]

    nama_pg = _split_f(f.get("pegawai_nama", ""))
    jawatan_pg = _split_f(f.get("pegawai_jawatan", ""))
    pg_rows = []
    for i, nm in enumerate(nama_pg):
        jw = jawatan_pg[i] if i < len(jawatan_pg) else ""
        pg_rows.append(f"{i+1}) {nm}" + (f"\n   {jw}" if jw else ""))
    _add_full_row("\n".join(pg_rows), bold=False, center=False, min_height="800")

    # Row 7: OBJEKTIF header + numbered list
    _add_full_row("OBJEKTIF", bold=True)
    objektif = f.get("objektif", "")
    if isinstance(objektif, list):
        obj_items = objektif
    else:
        obj_items = [o.strip() for o in objektif.split(",") if o.strip()]
    objektif_text = "\n".join(f"{i+1}) {o}" for i, o in enumerate(obj_items)) if obj_items else objektif
    _add_full_row(objektif_text, bold=False, center=False, min_height="800")

    # Row 8: RUMUSAN/LAPORAN (title row + content row)
    _add_full_row("RUMUSAN / LAPORAN", bold=True, center=False)
    _add_full_row(f.get("rumusan", ""), bold=False, center=False, min_height="2000")

    # Row 9: CADANGAN/TINDAKAN (title row + content row)
    _add_full_row("CADANGAN / TINDAKAN", bold=True, center=False)
    _add_full_row(f.get("cadangan", ""), bold=False, center=False, min_height="2000")

    # Row 10: DISEDIAKAN OLEH | DISAHKAN OLEH header
    _add_two_col_row("DISEDIAKAN OLEH", "DISAHKAN OLEH", left_bold=True, right_bold=True, right_center=True, min_height=None)

    # Row 11: Signature blocks side by side
    sig_left = (
        f"............................................................\n"
        f"NAMA    : {f.get('penyedia_nama', '')}\n"
        f"JAWATAN : {f.get('penyedia_jawatan', '')}\n"
        f"TARIKH  : {f.get('tarikh_disediakan', '')}"
    )
    sig_right = (
        f"............................................................\n"
        f"NAMA    : {f.get('pengesah_nama', '')}\n"
        f"JAWATAN : {f.get('pengesah_jawatan', '')}"
    )
    _add_two_col_row(sig_left, sig_right, left_bold=False, min_height="1600")

    # Set column widths
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = width

    # ── Lampiran Gambar ──
    images = get_report_images(session_id)
    if images:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(0)
        lh_p = doc.add_paragraph()
        lh_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lh_p.paragraph_format.space_before = Pt(0)
        lh_p.paragraph_format.space_after = Pt(6)
        lr = lh_p.add_run("LAMPIRAN GAMBAR")
        lr.font.name = "Arial"; lr.font.size = Pt(11); lr.bold = True

        # Fixed image dimensions — all same size, 2 per row
        IMG_W = Cm(7.8)
        IMG_H = Cm(5.2)  # ~landscape ratio 3:2

        for pair_start in range(0, len(images), 2):
            pair = images[pair_start:pair_start + 2]
            img_tbl = doc.add_table(rows=1, cols=2)
            img_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            img_tbl.autofit = False
            for j, img_info in enumerate(pair):
                cell = img_tbl.rows[0].cells[j]
                cell.width = Cm(8.5)
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                try:
                    from PIL import Image as PILImage
                    with PILImage.open(img_info["path"]) as im:
                        w, h = im.size
                    # Crop to uniform ratio before embedding — keep width fixed, calc height
                    run = para.add_run()
                    run.add_picture(img_info["path"], width=IMG_W, height=IMG_H)
                except Exception:
                    try:
                        para.add_run().add_picture(img_info["path"], width=IMG_W)
                    except Exception:
                        para.add_run(f"[Gambar {pair_start + j + 1}]")
            if len(pair) == 1:
                img_tbl.rows[0].cells[1].width = Cm(8.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def send_email(session_id: str, to_email: str, subject: str) -> dict:
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    document = get_document(session_id)
    if not document:
        return {"ok": False, "error": "Tiada laporan yang sedia untuk dihantar."}

    sender_email = os.getenv("SENDER_EMAIL")
    sender_password = os.getenv("SENDER_APP_PASSWORD")
    if not sender_email or not sender_password:
        return {"ok": False, "error": "Konfigurasi emel belum ditetapkan."}

    try:
        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = to_email
        msg["Subject"] = subject
        msg.attach(MIMEText(document, "plain", "utf-8"))

        filename = f"laporan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        attachment = MIMEBase("application", "octet-stream")
        attachment.set_payload(document.encode("utf-8"))
        encoders.encode_base64(attachment)
        attachment.add_header("Content-Disposition", f"attachment; filename={filename}")
        msg.attach(attachment)

        smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
        smtp_port = int(os.getenv("SMTP_PORT", "587"))
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)

        return {"ok": True, "message": f"Emel berjaya dihantar ke {to_email}."}
    except Exception as e:
        return {"ok": False, "error": f"Gagal menghantar emel: {e}"}


def clear_session(session_id: str):
    from backend.session_store import get_store
    get_store().delete_ns(session_id, _NS)
    clear_report_images(session_id)


def _try_parse_json(text: str) -> dict | None:
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        lines = lines[1:] if lines[0].startswith("```") else lines
        end = next((i for i, l in enumerate(lines) if l.strip() == "```"), len(lines))
        text = "\n".join(lines[:end])
    try:
        return json.loads(text)
    except (json.JSONDecodeError, ValueError):
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end > start:
            try:
                return json.loads(text[start:end + 1])
            except (json.JSONDecodeError, ValueError):
                pass
    return None
