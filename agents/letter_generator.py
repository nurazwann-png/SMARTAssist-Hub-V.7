"""Letter Generator Agent — Section 2: Personal Assistant for Official Correspondence."""

import json
import os
import re
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime
from backend.deepseek_client import chat_completion

# ── Document field schemas ──

SURAT_FIELDS = {
    "doc_type": "surat",
    "fields": [
        {"key": "rujukan", "label": "Ruj. Kami (Nombor Rujukan)", "example": "PPD.XXX/XXX/XX/XX ( )"},
        {"key": "tarikh", "label": "Tarikh", "example": "10 Julai 2026"},
        {"key": "penerima", "label": "Penerima", "example": "Semua Pengetua / Rujuk Senarai Edaran / SK Taman Maju"},
        {"key": "penerima_organisasi", "label": "Nama Organisasi Penerima", "example": "Jabatan Pendidikan Negeri Selangor", "optional": True},
        {"key": "penerima_alamat", "label": "Alamat Penerima", "example": "Aras 5, Blok E, 40000 Shah Alam, Selangor", "optional": True},
        {"key": "senarai_edaran", "label": "Senarai Edaran", "example": "[{\"nama\":\"\",\"jawatan\":\"Pengetua\",\"jabatan\":\"SK Taman Maju\"}]", "optional": True},
        {"key": "tajuk", "label": "Perkara / Tajuk Surat (huruf besar)", "example": "PERMOHONAN PERUNTUKAN KHAS"},
        {"key": "isi", "label": "Isi Kandungan Utama (pisahkan perenggan dengan baris kosong)", "example": "Penerangan tujuan surat.\n\nButiran lanjut."},
        {"key": "penandatangan_nama", "label": "Nama Penandatangan", "example": "Ahmad bin Ali"},
        {"key": "penandatangan_jawatan", "label": "Jawatan Penandatangan", "example": "Pegawai Pendidikan Daerah"},
        {"key": "pegawai_dihubungi", "label": "Pegawai Untuk Dihubungi", "example": "Puan Farah binti Zainudin, 084-123456", "optional": True},
        {"key": "nama_pejabat", "label": "Nama Pejabat / Unit", "example": "Pejabat Pendidikan Daerah Dalat"},
        {"key": "nama_organisasi", "label": "Nama Organisasi Penandatangan", "example": "Kementerian Pendidikan Malaysia", "optional": True},
        {"key": "salinan_kepada", "label": "Salinan Kepada (s.k.)", "example": "Pengarah JPN, Ketua Unit ICT", "optional": True},
    ],
}

MEMO_FIELDS = {
    "doc_type": "memo",
    "fields": [
        {"key": "rujukan", "label": "Nombor Rujukan (Ruj. Kami)", "example": "PPD.XXX-X/X/X ( )"},
        {"key": "tarikh", "label": "Tarikh Memo", "example": "10 Julai 2026"},
        {"key": "pengerusi", "label": "Nama Pengerusi dan Jawatan", "example": "Ahmad bin Ali (Pengetua SK Taman Jaya)"},
        {"key": "ahli", "label": "Nama Ahli-Ahli (pisahkan dengan koma)", "example": "Razif bin Ramli (Unit ICT), Nora binti Aziz (Unit HEM)"},
        {"key": "urus_setia", "label": "Nama Urus Setia dan Jawatan", "example": "Farah binti Zainudin (Pembantu Tadbir)"},
        {"key": "tajuk", "label": "Perkara / Tajuk Memo (huruf besar)", "example": "JEMPUTAN MESYUARAT PENGURUSAN BIL. 3/2026"},
        {"key": "tarikh_acara", "label": "Tarikh Acara", "example": "15 Julai 2026 (Isnin)"},
        {"key": "masa_acara", "label": "Masa Acara", "example": "8.00 pagi - 1.00 tengah hari"},
        {"key": "tempat_acara", "label": "Tempat Acara", "example": "Bilik Mesyuarat PPD Petaling Perdana"},
        {"key": "isi", "label": "Isi Kandungan", "example": "Butiran memo"},
        {"key": "langkah_kerja", "label": "Langkah Kerja / Tindakan (pisahkan dengan baris baru)", "example": "Semak senarai hadir\nSediakan kertas kerja", "optional": True},
        {"key": "penandatangan_nama", "label": "Nama Penandatangan", "example": "Ahmad bin Ali"},
        {"key": "penandatangan_jawatan", "label": "Jawatan Penandatangan", "example": "Pegawai Pendidikan Daerah"},
        {"key": "nama_pejabat", "label": "Nama Pejabat", "example": "Pejabat Pendidikan Daerah Petaling Perdana"},
    ],
}

SHARED_KEYS = {"rujukan", "tarikh", "tajuk", "isi"}

_SYSTEM_PROMPT_TEMPLATE = """Anda ialah pembantu peribadi profesional dalam SMARTAssist Hub.
Tugas anda: membantu pengguna menulis surat rasmi dan memo dalaman mengikut format KPM/Pekeliling Am.

CARA KERJA:
1. Kenal pasti jenis dokumen (surat rasmi atau memo dalaman)
2. Tanya maklumat yang diperlukan SATU DEMI SATU — HANYA SATU SOALAN setiap giliran
3. JANGAN terima semua maklumat sekaligus walaupun pengguna memberikan banyak maklumat dalam satu mesej — proses maklumat yang diberi, simpan dalam fields_collected, kemudian tanya field seterusnya yang belum diisi
4. Gunakan maklumat yang diberi untuk menjana dokumen lengkap apabila semua field dipenuhi
5. Pastikan tiada [PLACEHOLDER] yang belum diisi sebelum dokumen siap

URUTAN SOALAN (ikut susunan ini):
- Surat: rujukan → tarikh → penerima → (penerima_organisasi + penerima_alamat HANYA jika penerima adalah nama individu/organisasi — LANGKAU jika "Rujuk Senarai Edaran" atau "Semua...") → tajuk → (jana isi secara automatik) → penandatangan_nama → penandatangan_jawatan → nama_pejabat (pilihan) → nama_organisasi (pilihan) → salinan_kepada (pilihan)
- Jika penerima = "Senarai Edaran" atau "Rujuk Senarai Edaran": sistem akan tanya senarai_edaran berasingan melalui borang khas — JANGAN tanya pengguna untuk senarai_edaran dalam chat. JANGAN tanya penerima_organisasi dan penerima_alamat.
- Memo: rujukan → tarikh → pengerusi → ahli → urus_setia → tajuk → tarikh_acara → masa_acara → tempat_acara → (jana isi secara automatik) → langkah_kerja (tanya jika relevan) → penandatangan_nama → penandatangan_jawatan → nama_pejabat

GAYA:
- Bertanya seperti pembantu peribadi yang cekap — sopan, ringkas, dan spesifik
- HANYA tanya SATU field pada satu masa. Contoh: "Apakah nombor rujukan surat ini?"
- Jika pengguna beri beberapa maklumat sekaligus, terima semua yang diberi, kemudian tanya field seterusnya yang masih kosong
- Gunakan Bahasa Malaysia rasmi
- Format dokumen mengikut konvensyen rasmi KPM

TARIKH SEMASA: {current_date}
- Gunakan tarikh semasa sebagai rujukan. Tahun semasa ialah {current_year}.
- Jangan persoalkan tarikh yang pengguna berikan selagi formatnya betul.

FORMAT OUTPUT — balas HANYA dalam JSON:
{
  "phase": 0-4,
  "doc_type": "surat|memo|null",
  "message": "Mesej kepada pengguna",
  "field_asking": "nama field yang sedang ditanya (atau null)",
  "fields_collected": {"key": "value", ...},
  "document_preview": "Teks dokumen lengkap (hanya pada phase 3-4)",
  "validation_errors": ["Senarai ralat jika ada"],
  "ready_to_save": true/false
}

PHASES:
- Phase 0: Kenal pasti jenis dokumen (surat/memo). Jika pengguna minta tukar jenis, pindahkan field yang sama.
- Phase 1: Kumpul maklumat secara berperingkat (satu field setiap giliran). JANGAN tanya "isi" — isi akan dijana automatik.
- Phase 2: JANA ISI KANDUNGAN SECARA AUTOMATIK berdasarkan tajuk dan semua maklumat yang dikumpul. Simpan hasil dalam fields_collected dengan key "isi". Jika fields_collected sudah mengandungi "isi_user" (ringkasan/panduan dari pengguna melalui borang), WAJIB gunakan ia sebagai konteks untuk jana isi yang penuh, formal dan berformat dengan betul — JANGAN salin teks "isi_user" secara verbatim, jana semula dalam ayat rasmi yang lengkap. Jika maklumat tidak mencukupi, tanya soalan spesifik untuk mendapat konteks tambahan. JANGAN minta pengguna tulis isi sendiri.
  * Untuk SURAT: tulis isi dengan bernombor perenggan (2., 3., 4. dst), bahasa formal, lengkap dan profesional.
  * Untuk MEMO: field 'isi' WAJIB mengandungi SATU AYAT PENDEK SAHAJA tanpa sebarang newline — contoh: "Sukacita dimaklumkan bahawa mesyuarat akan diadakan seperti butiran berikut:". DILARANG KERAS memasukkan tarikh/masa/tempat, nombor perenggan (3., 4.), atau kandungan lain dalam 'isi'. Sistem akan papar tarikh_acara/masa_acara/tempat_acara secara berasingan. Field 'langkah_kerja' (PILIHAN) mengandungi langkah-langkah tindakan yang perlu diambil, SATU LANGKAH SETIAP BARIS (pisahkan dengan \n), contoh: "Semak senarai hadir\nSediakan kertas kerja\nHubungi peserta yang tidak hadir". Tanya tentang langkah_kerja hanya jika konteks memo memerlukan tindakan susulan.
- Phase 3: Tunjukkan pratonton dokumen lengkap — SEMAK tiada [PLACEHOLDER] kekal
- Phase 4: Dokumen disahkan dan sedia untuk dimuat turun / dihantar emel

FIELD KEYS YANG WAJIB DIGUNAKAN (guna key TEPAT ini dalam fields_collected — JANGAN guna nama lain):
Untuk surat: rujukan, tarikh, penerima, penerima_organisasi, penerima_alamat, tajuk, isi, penandatangan_nama, penandatangan_jawatan, salinan_kepada, senarai_edaran (pilihan)
PENTING: Gunakan "penerima" (BUKAN "penerima_nama"), simpan data senarai edaran dalam key "senarai_edaran"
- Field "penerima" boleh mengandungi sebarang teks: nama orang, jawatan, "Rujuk Senarai Edaran", "SK/SMK", dsb.
- PERATURAN SENARAI EDARAN — WAJIB DIIKUTI:
  * Jika penerima mengandungi perkataan "senarai edaran" (tidak kira huruf besar/kecil):
    a) Dalam blok alamat penerima, tulis HANYA: "Rujuk Senarai Edaran" (JANGAN tulis nama/jawatan individu)
    b) Di HUJUNG surat (selepas penandatangan), tambah halaman senarai edaran (lihat FORMAT SENARAI EDARAN)
  * Jika field "senarai_edaran" ada data JSON, gunakan data tersebut untuk jana halaman senarai edaran
  * JANGAN abaikan senarai_edaran — ia WAJIB dipaparkan sebagai halaman berasingan
Untuk memo: rujukan, tarikh, pengerusi, ahli, urus_setia, tajuk, tarikh_acara, masa_acara, tempat_acara, isi, langkah_kerja (pilihan), penandatangan_nama, penandatangan_jawatan, nama_pejabat

FORMAT TEMPLATE YANG MESTI DIIKUTI:

=== SURAT RASMI ===
                                      No. Rujukan : [rujukan]
                                      Tarikh      : [tarikh]

[penerima]          ← Jika senarai edaran: tulis "Rujuk Senarai Edaran" sahaja di sini
[penerima_organisasi]
[penerima_alamat]

[Panggilan Hormat — contoh: Tuan/Puan, atau Yang Berbahagia Dato',]

[TAJUK SURAT DALAM HURUF BESAR — RATA KIRI]

    Dengan hormatnya perkara di atas adalah dirujuk.

2.  [perenggan isi pertama]
3.  [perenggan isi kedua jika ada]
4.  [perenggan penutup]

Sekian, terima kasih.

"MALAYSIA MADANI"
"BERKHIDMAT UNTUK NEGARA"

Saya yang menjalankan amanah,


([NAMA PENANDATANGAN])
[Jawatan]

s.k.:
1. [salinan kepada jika ada]

=== MEMO DALAMAN ===
Kepada    | Pengerusi  : [pengerusi]
          | Ahli       : [ahli 1]
          |             : [ahli 2 dan seterusnya]
Daripada  | Urus setia : [urus_setia]
Tarikh               : [tarikh]
Perkara              : [TAJUK DALAM HURUF BESAR]
Ruj. Kami            : [rujukan]

Tuan,

Dengan segala hormatnya saya diarah merujuk kepada perkara di atas.

2.  [isi kandungan]

    Tarikh  : [tarikh_acara]
    Masa    : [masa_acara]
    Tempat  : [tempat_acara]

3.  [perenggan penutup]

Sekian, terima kasih

"MALAYSIA MADANI"
"BERKHIDMAT UNTUK NEGARA"

Saya yang menjalankan amanah


([NAMA PENANDATANGAN])
[Jawatan]
[Nama Pejabat]

=== FORMAT SENARAI EDARAN (halaman berasingan) ===
Jika field "senarai_edaran" mengandungi data (JSON array atau teks berformat), tambah halaman baru di hujung surat dengan format ini:

---

SENARAI EDARAN

BIL.  NAMA                        JAWATAN                     JABATAN/SEKOLAH
----  --------------------------  --------------------------  ---------------------------
1.    [NAMA PENERIMA]             [JAWATAN]                   [JABATAN/SEKOLAH]
2.    [NAMA PENERIMA]             [JAWATAN]                   [JABATAN/SEKOLAH]

PERATURAN SENARAI EDARAN:
- Semua teks dalam senarai edaran MESTI dalam HURUF BESAR (ALL CAPS) dan BOLD
- Jika nama kosong, biarkan ruang nama kosong — hanya jawatan dan jabatan
- Nombor BIL. berurutan dari 1
- Pisahkan halaman senarai edaran dengan "---" (garis pemisah) dan nyatakan "SENARAI EDARAN" sebagai tajuk bahagian

PENTING:
- Jangan reka maklumat PERIBADI (nama, jawatan, alamat) — tanya pengguna
- ISI KANDUNGAN surat/memo WAJIB dijana secara automatik oleh anda berdasarkan tajuk dan konteks. JANGAN tanya pengguna untuk menulis isi.
- Jika tajuk terlalu umum dan anda perlukan konteks tambahan, tanya soalan spesifik (contoh: "Apakah jumlah peruntukan yang dipohon?" atau "Berapa buah sekolah yang terlibat?")
- Gunakan KEY TEPAT seperti senarai di atas dalam fields_collected (contoh: "penerima", BUKAN "penerima_nama")
- Jika pengguna minta tukar dari surat ke memo (atau sebaliknya), pindahkan field yang sama (rujukan, tarikh, tajuk, isi)
- Pastikan JSON sah"""

_NS = "letter"  # namespace dalam SessionStore


def _get_session(session_id: str) -> dict:
    from backend.session_store import get_store
    store = get_store()
    data = store.get_all(session_id, _NS)
    if not data or "phase" not in data:
        default = {"phase": 0, "doc_type": None, "fields": {}, "document": None}
        store.set_all(session_id, _NS, default)
        return default
    return data


def _save_session(session_id: str, session: dict):
    from backend.session_store import get_store
    get_store().set_all(session_id, _NS, session)


def inject_pdf_context(session_id: str, fields: dict, doc_type: str) -> None:
    """Pra-isi session dengan maklumat yang diekstrak dari PDF.
    Dipanggil oleh /api/letter/upload-pdf sebelum pengguna hantar sebarang mesej."""
    session = _get_session(session_id)
    # Reset ke phase 0 supaya AI mula dari awal dengan konteks PDF
    session["phase"] = 0
    # Set doc_type cadangan jika valid
    if doc_type in ("surat", "memo"):
        session["doc_type"] = doc_type
    # Pra-isi hanya field yang ada nilai (tidak overwrite dengan null/kosong)
    # Map isi_pemakluman → isi field
    field_map = dict(fields)
    if "isi_pemakluman" in field_map:
        field_map["isi"] = field_map.pop("isi_pemakluman")

    # Simpan rujukan & tarikh dari PDF sebagai rujukan_asal/tarikh_asal (untuk ayat pemakluman)
    # Kosongkan rujukan supaya pengguna isi sendiri Ruj. Kami mereka
    rujukan_asal = str(field_map.get("rujukan", "") or "").strip()
    tarikh_asal  = str(field_map.get("tarikh", "") or "").strip()
    if rujukan_asal:
        session["fields"]["rujukan_asal"] = rujukan_asal
    if tarikh_asal:
        session["fields"]["tarikh_asal"] = tarikh_asal
    # Jangan pra-isi rujukan & tarikh — pengguna perlu isi Ruj. Kami & tarikh surat baharu mereka
    field_map.pop("rujukan", None)
    field_map.pop("tarikh", None)

    for k, v in field_map.items():
        if v and str(v).strip():
            session["fields"][k] = str(v).strip()

    # Bina ayat wajib pemakluman dan prepend ke isi
    ref_part  = f"No. Ruj: {rujukan_asal}" if rujukan_asal else "surat tersebut"
    date_part = f", bertarikh {tarikh_asal}" if tarikh_asal else ""
    ayat_wajib = f"Merujuk kepada surat {ref_part}{date_part}, adalah dimaklumkan perkara berikut untuk makluman dan tindakan pihak tuan/puan."

    # Fallback: jana isi asas dari tajuk jika isi_pemakluman tiada
    if not session["fields"].get("isi"):
        tajuk = session["fields"].get("tajuk", "")
        fallback_isi = (
            f"Adalah dimaklumkan bahawa pihak kami telah menerima surat berkenaan "
            f"{tajuk.replace('PEMAKLUMAN: ', '') if tajuk else 'perkara di atas'}. "
            f"Sehubungan dengan itu, perkara ini dimaklumkan kepada pihak tuan/puan untuk "
            f"makluman dan tindakan selanjutnya."
        )
        session["fields"]["isi"] = fallback_isi

    # Prepend ayat wajib ke isi (hanya jika belum ada)
    isi_semasa = session["fields"].get("isi", "")
    if "Merujuk kepada surat" not in isi_semasa:
        session["fields"]["isi"] = ayat_wajib + "\n\n" + isi_semasa
    # Flag untuk beritahu AI bahawa ada konteks dari PDF (surat pemakluman)
    session["pdf_context"] = True
    session["is_pemakluman"] = True
    _save_session(session_id, session)


def _get_field_schema(doc_type: str) -> list[dict]:
    if doc_type == "memo":
        return MEMO_FIELDS["fields"]
    return SURAT_FIELDS["fields"]


def _find_missing_fields(doc_type: str, collected: dict) -> list[dict]:
    schema = _get_field_schema(doc_type)
    base = [f for f in schema if not f.get("optional") and (f["key"] not in collected or not collected[f["key"]])]
    # When penerima is missing, inject penerima_organisasi/penerima_alamat right after it
    # so the form renders them as hideable fields (JS shows/hides based on penerima value)
    if doc_type == "surat" and any(f["key"] == "penerima" for f in base):
        extras = [f for f in schema if f["key"] in ("penerima_organisasi", "penerima_alamat")
                  and not collected.get(f["key"])]
        if extras:
            result = []
            for f in base:
                result.append(f)
                if f["key"] == "penerima":
                    result.extend(extras)
            return result
    return base


def _has_placeholders(text: str) -> list[str]:
    patterns = [
        r'\[PLACEHOLDER\]',
        r'\{\{CREATIVE:[^}]+\}\}',
        r'\{\{DYNAMIC_LIST:[^}]+\}\}',
    ]
    found = []
    for p in patterns:
        found.extend(re.findall(p, text))
    return found


def _build_document(doc_type: str, fields: dict) -> str:
    if doc_type == "memo":
        return _build_memo(fields)
    return _build_surat(fields)


import re as _re

def _strip_para_num(text: str) -> str:
    """Remove leading paragraph numbers like '2.', '3. ', '2)\t' from AI-generated isi."""
    return _re.sub(r'^\s*\d+[\.\)]\s*', '', text)


# Matches sub-item prefixes: 3.1  3.1.2  I.  II.  a.  a)  A)
_SUBITEM_RE = _re.compile(
    r'^(\d+(?:\.\d+)+\.?\s+'
    r'|[IVXivx]{1,5}\.\s+'
    r'|[a-zA-Z][.)]\s+)',
    _re.UNICODE
)


def _split_isi_lines(para_text: str):
    """Split isi paragraph into list of (prefix, body, is_subitem) tuples."""
    result = []
    for raw in para_text.split('\n'):
        line = raw.strip()
        if not line:
            continue
        m = _SUBITEM_RE.match(line)
        if m:
            pfx = line[:m.end()].rstrip()
            body = line[m.end():]
            result.append((pfx, body, True))
        else:
            result.append((None, line, False))
    return result


def _auto_panggilan(nama: str) -> str:
    n = nama.upper()
    if any(t in n for t in ["TAN SRI", "TUN "]):
        return "Yang Amat Berbahagia"
    if any(t in n for t in ["DATO'", "DATO ", "DATIN", "YBHG", "YBHg."]):
        return "Yang Berbahagia"
    if n.startswith("YB ") or " YB " in n:
        return "Yang Berhormat"
    if "PROF." in n or "PROFESOR" in n:
        return "Yang Dihormati Prof."
    if "DR." in n or " DR " in n:
        return "Yang Dihormati Dr."
    if any(t in n for t in ["ENCIK ", "EN. ", "EN."]) or n.startswith("ENCIK") or n.startswith("EN."):
        return "Tuan"
    if any(t in n for t in ["PUAN ", "CIK ", "DATIN"]) or n.startswith("PUAN") or n.startswith("CIK"):
        return "Puan"
    return "Tuan/Puan"


def _build_senarai_edaran_page(se_raw) -> str:
    """Jana halaman senarai edaran dari JSON string atau teks."""
    se_raw = str(se_raw) if se_raw else ""
    entries = []
    if se_raw and se_raw.strip():
        try:
            data = json.loads(se_raw)
            if isinstance(data, list):
                entries = data
        except Exception:
            # Fallback: baris per baris
            for line in se_raw.strip().splitlines():
                line = line.strip()
                if line:
                    entries.append({"nama": "", "jawatan": "", "jabatan": line})

    if not entries:
        return ""

    lines = [
        "",
        "---",
        "",
        "SENARAI EDARAN",
        "",
        f"{'BIL.':<6}{'NAMA':<30}{'JAWATAN':<30}JABATAN/SEKOLAH",
        f"{'----':<6}{'----------------------------':<30}{'----------------------------':<30}{'----------------------------'}",
    ]
    for i, entry in enumerate(entries, 1):
        nama    = str(entry.get("nama", "") or "").upper()
        jawatan = str(entry.get("jawatan", "") or "").upper()
        jabatan = str(entry.get("jabatan", "") or "").upper()
        lines.append(f"{str(i)+'.  ':<6}{nama:<30}{jawatan:<30}{jabatan}")

    return "\n".join(lines)


def _build_surat(f: dict) -> str:
    sk = str(f.get("salinan_kepada", "") or "")
    sk_lines = ""
    _sk_skip = {"tiada", "none", "kosong", "tak ada", "tidak ada", "-", "–", "tiada s.k.", "tiada sk"}
    if sk and sk.strip().lower() not in _sk_skip:
        items = [s.strip() for s in sk.split(",") if s.strip() and s.strip().lower() not in _sk_skip]
        if items:
            sk_lines = "\n\ns.k.:\n" + "\n".join(f"{i+1}. {item}" for i, item in enumerate(items))

    # Tentukan blok penerima — jika senarai edaran, tulis "Rujuk Senarai Edaran"
    penerima_raw = f.get('penerima', '') or f.get('penerima_nama', '[PLACEHOLDER]')
    is_se = 'senarai edaran' in penerima_raw.lower()
    penerima_line = "Rujuk Senarai Edaran" if is_se else penerima_raw

    # Ruj/Tarikh right-aligned block, then penerima address below
    RIGHT_W = 70
    addr_lines = [penerima_line]
    org = f.get('penerima_organisasi', '')
    alamat = f.get('penerima_alamat', '')
    if org and org != '[PLACEHOLDER]': addr_lines.append(org)
    if alamat and alamat != '[PLACEHOLDER]': addr_lines.append(alamat)

    header_lines = [
        f"{'No. Rujukan : ' + f.get('rujukan', '[PLACEHOLDER]'):>{RIGHT_W}}",
        f"{'Tarikh      : ' + f.get('tarikh', '[PLACEHOLDER]'):>{RIGHT_W}}",
        "",
    ] + addr_lines

    panggilan = "Tuan/Puan" if is_se else _auto_panggilan(penerima_raw)

    se_page = _build_senarai_edaran_page(f.get('senarai_edaran', '')) if is_se else ""

    return "\n".join(header_lines) + f"""

{panggilan},

{f.get('tajuk', '[PLACEHOLDER]').upper()}

    Dengan hormatnya perkara di atas adalah dirujuk.

{f.get('isi', '[PLACEHOLDER]')}

Sekian, terima kasih.

"MALAYSIA MADANI"
"BERKHIDMAT UNTUK NEGARA"

Saya yang menjalankan amanah,


({f.get('penandatangan_nama', '[PLACEHOLDER]').upper()})
{f.get('penandatangan_jawatan', '[PLACEHOLDER]')}""" + sk_lines + se_page


def _build_memo(f: dict) -> str:
    ahli_str = f.get('ahli', '[PLACEHOLDER]')
    if ahli_str and ahli_str != '[PLACEHOLDER]':
        ahli_list = [a.strip() for a in ahli_str.split(',') if a.strip()]
    else:
        ahli_list = ['[PLACEHOLDER]']

    # Col1=9chars, Col2=12chars — ensures all colons align at same position
    ahli_lines = f"          | {'Ahli':<12}: {ahli_list[0]}\n"
    for ahli in ahli_list[1:]:
        ahli_lines += f"          | {'':12}: {ahli}\n"

    return f"""MEMO DALAMAN

{'Kepada':<10}| {'Pengerusi':<12}: {f.get('pengerusi', '[PLACEHOLDER]')}
{ahli_lines}{'Daripada':<10}| {'Urus setia':<12}: {f.get('urus_setia', '[PLACEHOLDER]')}
{'Tarikh':<10}| {'':12}: {f.get('tarikh', '[PLACEHOLDER]')}
{'Perkara':<10}| {'':12}: {f.get('tajuk', '[PLACEHOLDER]').upper()}
{'Ruj. Kami':<10}| {'':12}: {f.get('rujukan', '[PLACEHOLDER]')}

Tuan,

Dengan segala hormatnya saya diarah merujuk kepada perkara di atas.

2.  {f.get('isi', '[PLACEHOLDER]')}

    {'Tarikh':<8}: {f.get('tarikh_acara', '[PLACEHOLDER]')}
    {'Masa':<8}: {f.get('masa_acara', '[PLACEHOLDER]')}
    {'Tempat':<8}: {f.get('tempat_acara', '[PLACEHOLDER]')}

3.  Kehadiran tuan/puan pada tarikh dan masa yang ditetapkan amatlah dihargai.

Sekian, terima kasih

"MALAYSIA MADANI"

"BERKHIDMAT UNTUK NEGARA"

Saya yang menjalankan amanah


({f.get('penandatangan_nama', '[PLACEHOLDER]').upper()})
{f.get('penandatangan_jawatan', '[PLACEHOLDER]')}
{f.get('nama_pejabat', '[PLACEHOLDER]')}"""


def _parse_form_submission(query: str, doc_type: str) -> dict:
    """Parse Label: value pairs from form submission messages directly into field keys.
    Returns dict of {field_key: value} for all matched labels."""
    schema = _get_field_schema(doc_type) if doc_type else SURAT_FIELDS["fields"] + MEMO_FIELDS["fields"]
    label_to_key = {f["label"].lower(): f["key"] for f in schema}
    result = {}
    import re as _re_form
    # Split on pattern: ". SomeLabel:" or ". CAPS word:"
    parts = _re_form.split(r'(?<=\.)\s+(?=[A-ZÀ-ɏ])', query)
    if len(parts) <= 1:
        parts = [query]
    for part in parts:
        m = _re_form.match(r'^(.+?):\s*(.+)$', part.strip(), _re_form.DOTALL)
        if not m:
            continue
        label_raw = m.group(1).strip().rstrip('.')
        value = m.group(2).strip().rstrip('.')
        if not value:
            continue
        # Match label to field key (case-insensitive)
        key = label_to_key.get(label_raw.lower())
        if key == 'isi':
            # Simpan sebagai isi_user — AI akan jana isi formal; jangan overwrite terus
            result['isi_user'] = value
        elif key:
            result[key] = value
    return result


def handle(query: str, history: list[dict] | None = None, session_id: str = "default", lang: str = "bm", user_name: str = "") -> str:
    sapaan = f", {user_name.split()[0]}" if user_name else ""
    if query == '__INTRO__':
        if lang == "en":
            return (f"Assalamualaikum and welcome{sapaan}! ✉️ I am the Official Letter Generator. I will help you prepare official letters, memos and circulars according to the correct KPM format. Could you tell me what type of letter you need and the basic information?\n\n"
                    "⚠️ Reminder: All documents generated by this system are AI-produced drafts. Please review and verify all content, names, dates and references carefully before official use.")
        return (f"Assalamualaikum dan selamat datang{sapaan}! ✉️ Saya Penjana Surat Rasmi. Saya akan membantu anda menyediakan surat rasmi, memo dan surat siaran mengikut format KPM yang betul. Boleh beritahu saya apakah jenis surat yang perlu disediakan dan maklumat asasnya?\n\n"
                "⚠️ Peringatan: Semua dokumen yang dijana adalah draf hasil AI. Sila semak dan sahkan semua kandungan, nama, tarikh serta rujukan dengan teliti sebelum digunakan secara rasmi.")

    session = _get_session(session_id)

    # Parse form submission directly — map all Label: value pairs to field keys
    # This bypasses the AI's one-field-at-a-time restriction for form submissions
    _form_parsed = _parse_form_submission(query, session.get("doc_type", ""))
    if _form_parsed:
        session["fields"].update({k: v for k, v in _form_parsed.items() if v})
        _save_session(session_id, session)

    _PATCH_KEYWORDS = {"kemaskini", "ubah", "ganti", "tukar", "edit", "update", "betulkan", "perbetulkan", "tukar kepada", "ubah suai"}
    query_lower = query.lower()
    if session.get("document") and any(kw in query_lower for kw in _PATCH_KEYWORDS):
        patch_prompt = f"""Dokumen semasa:
---
{session['document']}
---

Pengguna meminta perubahan berikut: {query}

Kembalikan HANYA dokumen yang telah dikemaskini dalam format JSON:
{{
  "phase": {session['phase']},
  "doc_type": "{session.get('doc_type', 'surat')}",
  "message": "Dokumen telah dikemaskini.",
  "document_preview": "<dokumen lengkap selepas perubahan>",
  "ready_to_save": true
}}"""
        now = datetime.now()
        date_str = now.strftime("%#d %B %Y") if os.name == "nt" else now.strftime("%-d %B %Y")
        system_prompt = _SYSTEM_PROMPT_TEMPLATE.replace("{current_date}", date_str).replace("{current_year}", str(now.year))
        _patch_lang_note = "\n\nIMPORTANT: The user has selected English. You MUST respond entirely in English for all 'message' fields." if lang == "en" else ""
        patch_messages = [{"role": "system", "content": system_prompt + _patch_lang_note}, {"role": "user", "content": patch_prompt}]
        try:
            raw = chat_completion(messages=patch_messages, temperature=0.2, max_tokens=3000)
            parsed = _try_parse_json(raw)
            if parsed and parsed.get("document_preview"):
                doc_text = parsed["document_preview"]
                session["document"] = doc_text
                parsed["ready_to_save"] = not _has_placeholders(doc_text)
                parsed["fields_status"] = {"collected": session["fields"], "missing": []}
                return json.dumps(parsed, ensure_ascii=False)
        except Exception:
            pass

    # Nota PDF context — hanya papar sekali, kemudian clear flag
    pdf_note = ""
    if session.get("pdf_context"):
        collected_fields = session.get("fields", {})
        has_isi = bool(collected_fields.get("isi"))
        pdf_note = f"""
ARAHAN PENTING — SURAT PEMAKLUMAN DARI PDF:
Pengguna telah memuat naik dokumen PDF rasmi. Sistem telah menganalisis dokumen tersebut dan mengekstrak maklumat berikut secara automatik (SUDAH DISIMPAN, tidak perlu tanya semula):
{json.dumps(collected_fields, ensure_ascii=False, indent=2)}

PERATURAN WAJIB — JANGAN LANGGAR:
1. Field "isi" TELAH DIJANA AUTOMATIK dari kandungan PDF. JANGAN tanya pengguna untuk isi kandungan surat. JANGAN minta pengguna tulis atau masukkan isi. Gunakan terus nilai "isi" yang telah ada.
2. Field "tajuk", "penerima", "penerima_organisasi" jika sudah ada — JANGAN tanya semula.
3. Ini adalah SURAT PEMAKLUMAN. Tajuk MESTI bermula dengan "PEMAKLUMAN:".
4. Field "rujukan" dan "tarikh" TIDAK ADA dalam fields_collected kerana pengguna perlu isi Ruj. Kami dan tarikh surat baharu mereka sendiri. WAJIB tanya kedua-dua field ini.

Field yang MASIH PERLU ditanya daripada pengguna (hanya yang belum ada nilai):
- Nombor rujukan surat baharu (rujukan) — Ruj. Kami pengguna, BUKAN rujukan surat asal PDF
- Tarikh surat baharu (tarikh) — tarikh surat baharu yang akan dijana, BUKAN tarikh surat asal PDF
- Penerima (penerima) — nama/jawatan/organisasi penerima jika belum ada
- Nama penandatangan (penandatangan_nama) — jika belum ada
- Jawatan penandatangan (penandatangan_jawatan) — jika belum ada
- Pegawai untuk dihubungi (pegawai_dihubungi) — nama dan nombor telefon pegawai hubungan

Sila:
a) Maklumkan ringkas kepada pengguna maklumat yang telah diekstrak dari PDF
b) Terus tanya HANYA maklumat yang masih belum lengkap (lihat senarai di atas)
c) JANGAN tanya isi kandungan surat — ia sudah dijana dari PDF
"""
        session["pdf_context"] = False
        _save_session(session_id, session)

    context_info = f"""
Status sesi semasa:
- Phase: {session['phase']}
- Jenis dokumen: {session['doc_type'] or 'belum ditentukan'}
- Field terkumpul: {json.dumps(session['fields'], ensure_ascii=False) if session['fields'] else 'tiada lagi'}
- Field belum diisi: {json.dumps([f['label'] for f in _find_missing_fields(session['doc_type'], session['fields'])], ensure_ascii=False) if session['doc_type'] else 'tentukan jenis dokumen dahulu'}
{pdf_note}"""

    now = datetime.now()
    date_str = now.strftime("%#d %B %Y") if os.name == "nt" else now.strftime("%-d %B %Y")
    system_prompt = _SYSTEM_PROMPT_TEMPLATE.replace("{current_date}", date_str).replace("{current_year}", str(now.year))
    lang_note = "\n\nIMPORTANT: The user has selected English. You MUST respond entirely in English. All 'message' and text fields in your JSON response must be in English. The generated document content should remain in Malay as it is an official KPM document." if lang == "en" else ""

    messages = [
        {"role": "system", "content": system_prompt + context_info + lang_note},
    ]
    if history:
        for msg in history[-10:]:
            messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": query})

    try:
        raw = chat_completion(messages=messages, temperature=0.3, max_tokens=3000)
    except RuntimeError as e:
        return json.dumps({
            "phase": session["phase"],
            "message": f"Ralat menghubungi API. Sila cuba lagi.\n{e}",
            "ready_to_save": False,
        }, ensure_ascii=False)

    parsed = _try_parse_json(raw)
    if not parsed:
        return json.dumps({
            "phase": session["phase"],
            "message": raw,
            "ready_to_save": False,
        }, ensure_ascii=False)

    if parsed.get("doc_type") and parsed["doc_type"] in ("surat", "memo"):
        old_type = session["doc_type"]
        new_type = parsed["doc_type"]
        if old_type and old_type != new_type and session["fields"]:
            carried = {k: v for k, v in session["fields"].items() if k in SHARED_KEYS}
            session["fields"] = carried
        session["doc_type"] = new_type

    if parsed.get("fields_collected"):
        fc = parsed["fields_collected"]
        # Normalize: map old key penerima_nama → penerima
        if "penerima_nama" in fc and "penerima" not in fc:
            fc["penerima"] = fc.pop("penerima_nama")
        if "penerima_jawatan" in fc:
            # If penerima exists, append jawatan to it; else discard (merged field)
            if fc.get("penerima") and fc["penerima_jawatan"]:
                fc["penerima"] = f"{fc['penerima']} ({fc.pop('penerima_jawatan')})"
            else:
                fc.pop("penerima_jawatan", None)
        session["fields"].update(fc)
        # For memo, isi must be one sentence only — trim anything after first newline
        if session.get("doc_type") == "memo" and "isi" in session["fields"]:
            isi_val = session["fields"]["isi"]
            if "\n" in isi_val:
                session["fields"]["isi"] = isi_val.split("\n")[0].strip()

    # Extract senarai_edaran from raw query if not yet captured (AI may miss it)
    if "senarai_edaran" not in session["fields"]:
        import re as _re2
        se_match = _re2.search(r'[Ss]enarai [Ee]daran\s*:\s*(\[.*?\])', query, _re2.DOTALL)
        if se_match:
            session["fields"]["senarai_edaran"] = se_match.group(1).strip()

    # Extract Isi Kandungan from form message reliably (LLM may ignore it per system prompt)
    import re as _re3
    _isi_m = _re3.search(
        r'Isi Kandungan\s*:\s*(.+?)(?=\s*,\s*[A-Z][a-zA-Z\s]*\s*:|\s*$)',
        query, _re3.DOTALL | _re3.IGNORECASE
    )
    if _isi_m:
        _user_isi = _isi_m.group(1).strip().strip(',').strip()
        if _user_isi and "isi_user" not in session["fields"]:
            session["fields"]["isi_user"] = _user_isi

    if parsed.get("phase") is not None:
        session["phase"] = parsed["phase"]

    # For memo: build proper isi. If user provided a brief, incorporate it (override LLM version).
    if session.get("doc_type") == "memo":
        tajuk = session["fields"].get("tajuk", "")
        isi_user = session["fields"].get("isi_user", "").strip()
        if tajuk and isi_user:
            _brief = isi_user.rstrip('.').rstrip(',').strip()
            session["fields"]["isi"] = (
                f"Sukacita dimaklumkan bahawa {tajuk.rstrip('.')} diadakan bagi tujuan "
                f"{_brief[0].lower()}{_brief[1:]} seperti butiran berikut:"
            )
        elif tajuk and "isi" not in session["fields"]:
            required_keys = {f["key"] for f in MEMO_FIELDS["fields"] if not f.get("optional") and f["key"] != "isi"}
            if required_keys.issubset(session["fields"].keys()):
                session["fields"]["isi"] = f"Sukacita dimaklumkan bahawa {tajuk.rstrip('.')} akan diadakan seperti butiran berikut:"

    # Jana isi surat secara automatik jika isi_user ada tapi isi belum dijana
    if (session.get("doc_type") == "surat"
            and session["fields"].get("isi_user")
            and not session["fields"].get("isi")):
        _tajuk = session["fields"].get("tajuk", "")
        _isi_user = session["fields"].get("isi_user", "")
        _penerima = session["fields"].get("penerima", "")
        _isi_prompt = [
            {"role": "system", "content": (
                "Kamu adalah pembantu yang menulis isi kandungan surat rasmi dalam Bahasa Malaysia yang formal dan profesional. "
                "Balas HANYA dengan teks isi kandungan sahaja — tiada penjelasan, tiada JSON."
            )},
            {"role": "user", "content": (
                f"Jana isi kandungan surat rasmi berdasarkan maklumat berikut:\n"
                f"Tajuk: {_tajuk}\n"
                f"Penerima: {_penerima}\n"
                f"Ringkasan isi dari pengguna: {_isi_user}\n\n"
                "Tulis 2-4 perenggan bernombor (2., 3., 4. dst) yang formal, lengkap dan profesional. "
                "Setiap perenggan bermula dengan nombor (contoh: '2. Sehubungan dengan itu...'). "
                "Jangan masukkan 'Dengan segala hormatnya' atau 'Sekian'. "
                "Gunakan Bahasa Malaysia rasmi."
            )}
        ]
        try:
            _isi_raw = chat_completion(messages=_isi_prompt, temperature=0.4, max_tokens=800)
            if _isi_raw and _isi_raw.strip():
                session["fields"]["isi"] = _isi_raw.strip()
                _save_session(session_id, session)
        except Exception:
            # Fallback: guna isi_user terus jika API call gagal
            session["fields"]["isi"] = f"2. {_isi_user}\n\n3. Kerjasama dan perhatian pihak tuan/puan dalam perkara ini amatlah dihargai."
            _save_session(session_id, session)

    all_filled = session["doc_type"] and not _find_missing_fields(session["doc_type"], session["fields"])
    if session["doc_type"] and (session["phase"] >= 3 or all_filled):
        if all_filled:
            session["phase"] = max(session["phase"], 3)
            parsed["phase"] = session["phase"]
        doc_text = _build_document(session["doc_type"], session["fields"])
        placeholders = _has_placeholders(doc_text)
        parsed["document_preview"] = doc_text
        if session["doc_type"] == "memo":
            parsed["document_html"] = _build_memo_html(session["fields"])
        elif session["doc_type"] == "surat":
            parsed["document_html"] = _build_surat_html(session["fields"])
        if placeholders:
            parsed["validation_errors"] = [f"Masih ada placeholder yang belum diisi: {', '.join(placeholders)}"]
            parsed["ready_to_save"] = False
        else:
            parsed["ready_to_save"] = True
            session["document"] = doc_text

    # form_extras: optional fields that should appear in the form but don't block generation
    form_extras: list[str] = []
    if session.get("doc_type") == "surat":
        import re as _re
        penerima_val = str(session["fields"].get("penerima", "") or "").strip()
        is_specific_penerima = penerima_val and not _re.search(
            r'senarai\s*edaran|^semua|^rujuk|^\s*sk\b|^\s*smk\b', penerima_val, _re.IGNORECASE
        )
        if is_specific_penerima:
            surat_schema = _get_field_schema("surat")
            for f in surat_schema:
                if f["key"] in ("penerima_organisasi", "penerima_alamat") and not (session["fields"].get(f["key"])):
                    form_extras.append(f["label"])

    _missing_all = _find_missing_fields(session["doc_type"], session["fields"]) if session["doc_type"] else []
    # Jangan tunjuk 'isi' dalam borang lagi jika isi_user sudah ada (user dah isi sekali)
    _has_isi_user = bool(session["fields"].get("isi_user"))
    _missing_form = [f["label"] for f in _missing_all if not (f["key"] == "isi" and _has_isi_user)]
    parsed["fields_status"] = {
        "collected": session["fields"],
        "missing": _missing_form,
        "form_extras": form_extras,
    }

    _save_session(session_id, session)
    return json.dumps(parsed, ensure_ascii=False)


def get_document(session_id: str) -> str | None:
    session = _get_session(session_id)
    if session and session.get("document"):
        return session["document"]
    return None


def get_fields(session_id: str) -> dict:
    session = _get_session(session_id)
    return session["fields"].copy() if session else {}


def apply_improvement(session_id: str, improved_fields: dict) -> str | None:
    """Update session fields and rebuild document using original template. Returns new doc text."""
    session = _get_session(session_id)
    if not session or not session.get("doc_type"):
        return None
    for key in ("isi",):
        if key in improved_fields and improved_fields[key]:
            session["fields"][key] = improved_fields[key]
    # Memo isi must stay single sentence
    if session.get("doc_type") == "memo":
        isi = session["fields"].get("isi", "")
        if "\n" in isi:
            session["fields"]["isi"] = isi.split("\n")[0].strip()
    # Rebuild document from updated fields using original template
    new_doc = _build_document(session["doc_type"], session["fields"])
    session["document"] = new_doc
    _save_session(session_id, session)
    return new_doc


GENERATED_FIELD_KEYS = ["isi"]


def build_docx(session_id: str) -> bytes | None:
    session = _get_session(session_id)
    doc_text = get_document(session_id)
    if not doc_text:
        return None

    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = DocxDocument()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Insert letterhead + page-continuation footer (surat only)
    _doc_type_check = session.get("doc_type", "surat") if session else "surat"
    try:
        from backend.letterhead_store import get_active_path_by_type
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlElement
        lh_path = get_active_path_by_type("letterhead") if _doc_type_check != "memo" else None
        if _doc_type_check != "memo":
            section = doc.sections[0]
            section.header_distance = Cm(0.5)
            # First-page header differs from subsequent pages
            section.different_first_page_header_footer = True

            if lh_path:
                # Page 1 header — letterhead + horizontal rule
                import io as _lh_io
                from PIL import Image as _lh_PIL
                _lh_img = _lh_PIL.open(str(lh_path)).convert("RGBA")
                _lh_buf = _lh_io.BytesIO()
                _lh_img.save(_lh_buf, format="PNG")
                _lh_buf.seek(0)
                fp_header = section.first_page_header
                for p in fp_header.paragraphs:
                    p.clear()
                lh_para = fp_header.paragraphs[0]
                lh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lh_para.paragraph_format.space_before = Pt(0)
                lh_para.paragraph_format.space_after = Pt(4)
                lh_para.add_run().add_picture(_lh_buf, width=Cm(17))
                hr_para = fp_header.add_paragraph()
                hr_para.paragraph_format.space_before = Pt(0)
                hr_para.paragraph_format.space_after = Pt(0)
                pPr = hr_para._p.get_or_add_pPr()
                pBdr = _OxmlElement('w:pBdr')
                bottom = _OxmlElement('w:bottom')
                bottom.set(_qn('w:val'), 'single')
                bottom.set(_qn('w:sz'), '6')
                bottom.set(_qn('w:space'), '1')
                bottom.set(_qn('w:color'), '000000')
                pBdr.append(bottom)
                pPr.append(pBdr)

            # Pages 2+ header — empty
            reg_header = section.header
            for p in reg_header.paragraphs:
                p.clear()

            def _r(*kids):
                r = _OxmlElement('w:r')
                for k in kids: r.append(k)
                return r

            def _rpr():
                rpr = _OxmlElement('w:rPr')
                for tag in ('w:sz', 'w:szCs'):
                    el = _OxmlElement(tag); el.set(_qn('w:val'), '20'); rpr.append(el)
                return rpr

            def _fc(typ, dirty=False):
                fc = _OxmlElement('w:fldChar'); fc.set(_qn('w:fldCharType'), typ)
                if dirty: fc.set(_qn('w:dirty'), 'true')
                return fc

            def _instr(s):
                it = _OxmlElement('w:instrText'); it.set(_qn('xml:space'), 'preserve'); it.text = s; return it

            def _t(s):
                t = _OxmlElement('w:t'); t.set(_qn('xml:space'), 'preserve'); t.text = s; return t

            # Estimate if document will be multi-page (>1 page needs footer)
            _sess_fields = session.get("fields", {}) if session else {}
            _isi_raw_est = _sess_fields.get('isi', '')
            _isi_paras_est = [_strip_para_num(p.strip()) for p in _isi_raw_est.split('\n\n') if p.strip()] if _isi_raw_est else []
            _total_isi_lines_est = sum(len(_split_isi_lines(p)) for p in _isi_paras_est)
            _is_multipage = _total_isi_lines_est > 8

            # Page 1 footer: ..{PAGE+1}/- right-aligned — only for multi-page docs
            if _is_multipage:
                fp_ftr = section.first_page_footer
                fp_ftr.is_linked_to_previous = False
                para1 = fp_ftr.paragraphs[0]
                para1.clear()
                para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para1.paragraph_format.space_before = Pt(0)
                para1.paragraph_format.space_after = Pt(0)
                p1 = para1._p
                p1.append(_r(_rpr(), _t('..')))
                p1.append(_r(_fc('begin', dirty=True)))
                p1.append(_r(_instr(' = ')))
                p1.append(_r(_fc('begin'))); p1.append(_r(_instr(' PAGE ')))
                p1.append(_r(_fc('separate'))); p1.append(_r(_t('1'))); p1.append(_r(_fc('end')))
                p1.append(_r(_instr(r' + 1 \# "0" ')))
                p1.append(_r(_fc('separate'))); p1.append(_r(_t('2'))); p1.append(_r(_fc('end')))
                p1.append(_r(_rpr(), _t('/-')))

            # Pages 2+ footer: centred page number { PAGE } — only for multi-page docs
            if _is_multipage:
                reg_ftr = section.footer
                reg_ftr.is_linked_to_previous = False
                para2 = reg_ftr.paragraphs[0]
                para2.clear()
                para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para2.paragraph_format.space_before = Pt(0)
                para2.paragraph_format.space_after = Pt(0)
                p2 = para2._p
                p2.append(_r(_fc('begin', dirty=True)))
                p2.append(_r(_instr(' PAGE ')))
                p2.append(_r(_fc('separate'))); p2.append(_r(_t('2'))); p2.append(_r(_fc('end')))
    except Exception:
        pass

    doc_type = session.get("doc_type", "surat") if session else "surat"
    fields = session.get("fields", {}) if session else {}

    if doc_type == "memo":
        # Memo: logo in body (no letterhead header)
        try:
            from backend.letterhead_store import get_active_path_by_type
            logo_path = get_active_path_by_type("logo")
            if logo_path:
                import io as _io
                from PIL import Image as _PILImage
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                # Convert to PNG in memory so python-docx can always read it
                _img = _PILImage.open(str(logo_path)).convert("RGBA")
                _buf = _io.BytesIO()
                _img.save(_buf, format="PNG")
                _buf.seek(0)
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_para.paragraph_format.space_before = Pt(0)
                logo_para.paragraph_format.space_after = Pt(6)
                logo_para.add_run().add_picture(_buf, width=Cm(7.67), height=Cm(4.31))
        except Exception:
            pass
        _build_memo_docx(doc, fields)
    else:
        _build_surat_docx(doc, doc_text, fields)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_surat_docx(doc, doc_text: str, fields: dict = None):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _no_border(tbl):
        tbl_element = tbl._tbl
        tbl_pr = tbl_element.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl_element.insert(0, tbl_pr)
        tbl_borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'none')
            tbl_borders.append(b)
        tbl_pr.append(tbl_borders)

    def _p(text="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, indent_cm=0, size=12):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        if indent_cm:
            para.paragraph_format.first_line_indent = Cm(indent_cm)
        if text:
            run = para.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(size)
            run.bold = bold
        return para

    if fields:
        # Ruj. Kami + Tarikh right-aligned (per KPM format)
        _p(f"Ruj. Kami : {fields.get('rujukan', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _p(f"Tarikh    : {fields.get('tarikh', '')}", align=WD_ALIGN_PARAGRAPH.RIGHT)

        doc.add_paragraph("")

        penerima_raw = fields.get('penerima', '') or fields.get('penerima_nama', '')
        is_se = 'senarai edaran' in penerima_raw.lower()
        penerima_line = "Rujuk Senarai Edaran" if is_se else penerima_raw
        for line in [
            penerima_line,
            fields.get('penerima_organisasi', ''),
            fields.get('penerima_alamat', ''),
        ]:
            if line:
                _p(line)

        doc.add_paragraph("")

        panggilan = "Tuan/Puan" if is_se else _auto_panggilan(penerima_raw)
        _p(f"{panggilan},")
        doc.add_paragraph("")

        # Tajuk — LEFT, bold, uppercase (per template)
        _p(fields.get('tajuk', '').upper(), bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        doc.add_paragraph("")

        # Paragraph 1 — no indent (per template line 43)
        _p("Dengan segala hormatnya perkara di atas adalah dirujuk.",
           align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        doc.add_paragraph("")

        # Isi paragraphs numbered from 2. with hanging indent
        isi_raw = fields.get('isi', '')
        isi_paras = [_strip_para_num(p.strip()) for p in isi_raw.split('\n\n') if p.strip()] if isi_raw else []

        def _add_isi_para(text, left_cm, hanging_cm, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Cm(left_cm)
            p.paragraph_format.first_line_indent = Cm(-hanging_cm)
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(12)
            return p

        total_isi_lines = sum(len(_split_isi_lines(p)) for p in isi_paras)

        for i, para_text in enumerate(isi_paras):
            lines = _split_isi_lines(para_text)
            if not lines:
                continue
            _, first_body, _ = lines[0]
            has_children = len(lines) > 1
            if has_children:
                # Has sub-items/continuations — use hanging indent
                _add_isi_para(f"{i+2}.\t{first_body}", left_cm=1.25, hanging_cm=1.25)
                for pfx, body, is_sub in lines[1:]:
                    if is_sub:
                        _add_isi_para(f"{pfx}\t{body}", left_cm=2.5, hanging_cm=1.25)
                    else:
                        _add_isi_para(body, left_cm=1.25, hanging_cm=0)
            else:
                # Plain paragraph — no hanging indent
                _add_isi_para(f"{i+2}.\t{first_body}", left_cm=0, hanging_cm=0)
            doc.add_paragraph("")

        _p("Sekian, terima kasih.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        doc.add_paragraph("")
        _p('"MALAYSIA MADANI"', bold=True)
        doc.add_paragraph("")
        _p('"BERKHIDMAT UNTUK NEGARA"', bold=True)
        doc.add_paragraph("")
        _p("Saya yang menjalankan amanah,")
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")
        _p(f"({fields.get('penandatangan_nama', '').upper()})", bold=True)
        _p(fields.get('penandatangan_jawatan', ''))
        if fields.get('nama_pejabat'):
            _p(fields.get('nama_pejabat', ''))
        if fields.get('nama_organisasi'):
            _p(fields.get('nama_organisasi', ''))

        _sk_skip_set = {"tiada", "none", "kosong", "tak ada", "tidak ada", "-", "–", "tiada s.k.", "tiada sk"}
        sk = fields.get('salinan_kepada', '')
        if sk and sk.strip().lower() not in _sk_skip_set:
            sk_items = [s.strip() for s in sk.split(',') if s.strip() and s.strip().lower() not in _sk_skip_set]
            if sk_items:
                doc.add_paragraph("")
                _p("s.k.:")
                for i, item in enumerate(sk_items):
                    _p(f"{i+1}. {item}")

        # Senarai edaran — halaman baru dalam DOCX
        if is_se:
            se_raw = fields.get('senarai_edaran', '')
            se_entries = []
            try:
                se_data = json.loads(se_raw) if se_raw else []
                if isinstance(se_data, list):
                    se_entries = se_data
            except Exception:
                pass
            if se_entries:
                doc.add_page_break()
                heading = doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = heading.add_run("SENARAI EDARAN")
                run.bold = True
                run.font.size = Pt(13)
                run.font.name = "Arial"
                doc.add_paragraph("")
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                for cell, txt in zip(hdr, ["BIL.", "NAMA", "JAWATAN", "JABATAN / SEKOLAH"]):
                    cell.text = txt
                    cell.paragraphs[0].runs[0].bold = True
                for i, entry in enumerate(se_entries, 1):
                    row = table.add_row().cells
                    row[0].text = f"{i}."
                    row[1].text = str(entry.get('nama', '') or '').upper()
                    row[2].text = str(entry.get('jawatan', '') or '').upper()
                    row[3].text = str(entry.get('jabatan', '') or '').upper()
                    for cell in row:
                        cell.paragraphs[0].runs[0].bold = True if cell.text else False
    else:
        # Fallback: plain line-by-line rendering
        lines = doc_text.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            is_bold = False
            if stripped.startswith('"') and stripped.endswith('"'):
                is_bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif stripped.startswith("s.k.:"):
                is_bold = True
            elif stripped == stripped.upper() and len(stripped) > 5 and stripped[0].isalpha():
                is_bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(stripped)
            run.font.size = Pt(11)
            run.font.name = "Arial"
            if is_bold:
                run.bold = True


def _build_memo_html(f: dict) -> str:
    logo_url = _get_lh_img_url("logo")
    if logo_url:
        logo_block = f'<img src="{logo_url}" style="max-width:100%;max-height:120px;display:block;margin:0 auto 8px auto">'
    else:
        org_name = f.get("nama_pejabat", "")
        logo_block = (
            f'<div style="text-align:center;padding:8px 0 4px 0">'
            f'<div style="font-size:12pt;font-weight:bold;line-height:1.8">{org_name}</div>'
            f'</div>'
        ) if org_name else ""

    ahli_str = f.get('ahli', '')
    ahli_list = [a.strip() for a in ahli_str.split(',') if a.strip()] if ahli_str else []
    pengerusi = f.get('pengerusi', '')

    TD1 = 'style="padding:4px 8px;border:1px solid #000;vertical-align:top;width:110px"'
    TD2 = 'style="padding:4px 6px;border:1px solid #000;vertical-align:top;width:18px;text-align:center"'
    TD3 = 'style="padding:4px 8px;border:1px solid #000;vertical-align:top"'
    rows = []
    if pengerusi:
        rows.append((f'<b>Kepada</b>', ':', pengerusi))
        for ahli in ahli_list:
            if ahli:
                rows.append(('', '', ahli))
    rows.append((f'<b>Daripada</b>', ':', f.get('urus_setia', '')))
    rows.append((f'<b>Tarikh</b>', ':', f.get('tarikh', '')))
    rows.append((f'<b>Ruj. Kami</b>', ':', f.get('rujukan', '')))
    rows.append((f'<b>Perkara</b>', ':', f'<b>{f.get("tajuk", "").upper()}</b>'))

    rows_html = ''.join(
        f'<tr>'
        f'<td {TD1}>{r[0]}</td>'
        f'<td {TD2}>{r[1]}</td>'
        f'<td {TD3}>{r[2]}</td>'
        f'</tr>'
        for r in rows
    )

    panggilan = _auto_panggilan(f.get('pengerusi', 'Tuan'))
    tarikh_acara = f.get('tarikh_acara', '')
    masa_acara = f.get('masa_acara', '')
    tempat_acara = f.get('tempat_acara', '')

    # Strip isi from AI-added tarikh/masa/tempat info (already in dedicated block)
    _isi_raw = _strip_para_num(str(f.get('isi', '') or '').split('\n')[0].strip())
    _isi_lower = _isi_raw.lower()
    # Remove trailing sentences that mention date/time/place
    for _kw in [tarikh_acara.lower(), masa_acara.lower(), tempat_acara.lower()]:
        if _kw and _kw in _isi_lower:
            # Truncate isi at the sentence that first mentions these details
            for _sent_end in ['. ', ', ']:
                _idx = _isi_lower.find(_kw)
                _start = _isi_raw.rfind('.', 0, _idx)
                if _start != -1:
                    _isi_raw = _isi_raw[:_start + 1].strip()
                    _isi_lower = _isi_raw.lower()
                    break
    isi = _isi_raw.replace('\n', '<br>')

    # Langkah kerja — numbered paragraphs starting at 4
    langkah_str = f.get('langkah_kerja', '')
    langkah_list = [l.strip() for l in langkah_str.split('\n') if l.strip()] if langkah_str else []

    # No hanging indent — continuation wraps under number (per user request)
    hang = 'style="margin:6px 0;line-height:1.6;text-align:justify"'
    acara_indent = 'style="margin:2px 0 2px 4em;line-height:1.6"'
    normal = 'style="margin:6px 0;line-height:1.6"'

    langkah_html = ''
    for i, step in enumerate(langkah_list):
        num = i + 4
        langkah_html += f'<p {hang}>{num}.&nbsp;&nbsp;&nbsp;&nbsp;{step}</p>'

    return (
        f'<div style="font-family:Arial,sans-serif;font-size:12pt;line-height:1.5;color:#000">'
        f'{logo_block}'
        f'<table style="width:100%;border-collapse:collapse;margin-bottom:12px">'
        f'<tr><td colspan="3" style="background:#000;color:#fff;text-align:center;font-weight:bold;'
        f'padding:6px 8px;font-size:12pt;letter-spacing:1px">MEMO DALAMAN</td></tr>'
        f'{rows_html}'
        f'</table>'
        f'<p {normal}>{panggilan},</p>'
        f'<p {normal}>Dengan segala hormatnya perkara di atas adalah dirujuk.</p>'
        f'<p {hang}>2.&nbsp;&nbsp;&nbsp;&nbsp;{isi}</p>'
        f'<p {acara_indent}><b>Tarikh</b>&emsp;&nbsp;: {tarikh_acara}</p>'
        f'<p {acara_indent}><b>Masa</b>&emsp;&emsp;: {masa_acara}</p>'
        f'<p {acara_indent}><b>Tempat</b>&emsp;: {tempat_acara}</p>'
        f'<p {hang}>3.&nbsp;&nbsp;&nbsp;&nbsp;Kehadiran tuan/puan pada tarikh dan masa yang ditetapkan amatlah dihargai.</p>'
        f'{langkah_html}'
        f'<p {normal}>Sekian, terima kasih.</p>'
        f'<div style="page-break-inside:avoid">'
        f'<p style="margin:11pt 0 0 0;line-height:1.6"><b>&ldquo;MALAYSIA MADANI&rdquo;</b></p>'
        f'<p style="margin:11pt 0 0 0;line-height:1.6"><b>&ldquo;BERKHIDMAT UNTUK NEGARA&rdquo;</b></p>'
        f'<p style="margin:11pt 0 0 0;line-height:1.6">Saya yang menjalankan amanah,</p>'
        f'<p style="margin:33pt 0 0 0;line-height:1.6"><b>({f.get("penandatangan_nama", "").upper()})</b></p>'
        f'<p {normal}>{f.get("penandatangan_jawatan", "")}</p>'
        f'<p {normal}>{f.get("nama_pejabat", "")}</p>'
        f'</div>'
        f'</div>'
    )


def _get_lh_img_url(lh_type: str = "letterhead") -> str | None:
    try:
        from backend.letterhead_store import get_active_by_type
        lh = get_active_by_type(lh_type)
        if lh and lh.get("filename"):
            return f"/api/letterhead/image/{lh['filename']}"
    except Exception:
        pass
    return None


def _build_surat_html(f: dict) -> str:
    lh_url = _get_lh_img_url("letterhead")
    if lh_url:
        lh_block = f'<img src="{lh_url}" style="max-width:100%;max-height:150px;display:block;margin:0 auto">'
    else:
        org_name = f.get("nama_pejabat", "")
        lh_block = (
            f'<div style="text-align:center;padding:12px 0 8px 0">'
            f'<div style="font-size:13pt;font-weight:bold;text-transform:uppercase;line-height:1.8">'
            f'{org_name or "NAMA ORGANISASI"}</div>'
            f'</div>'
        )

    sk = str(f.get("salinan_kepada", "") or "")
    sk_html = ""
    _sk_skip = {"tiada", "none", "kosong", "tak ada", "tidak ada", "-", "–", "tiada s.k.", "tiada sk"}
    if sk and sk.strip().lower() not in _sk_skip:
        items = [s.strip() for s in sk.split(",") if s.strip() and s.strip().lower() not in _sk_skip]
        if items:
            sk_html = '<p style="margin:16px 0 4px 0"><b>s.k.:</b></p>'
            for i, item in enumerate(items):
                sk_html += f'<p style="margin:2px 0">{i+1}. {item}</p>'

    # Penerima block — guna field 'penerima' baru, fallback ke penerima_nama lama
    penerima_raw = str(f.get('penerima', '') or f.get('penerima_nama', '') or '')
    is_se = 'senarai edaran' in penerima_raw.lower()
    penerima_line = "Rujuk Senarai Edaran" if is_se else penerima_raw
    panggilan = "Tuan/Puan" if is_se else _auto_panggilan(penerima_raw)

    isi_raw = str(f.get('isi', '') or '')
    import re as _re_isi
    _DENGAN_RE = _re_isi.compile(r'^\s*dengan\s+segala\s+hormatnya', _re_isi.IGNORECASE)
    _TARIKH_BARIS_RE = _re_isi.compile(r'^\s*(Tarikh|Masa|Tempat|Venue|Date|Time)\s*:', _re_isi.IGNORECASE)
    _plain_p = 'margin:6px 0;line-height:1.6;text-align:justify'
    _indent_p = 'margin:2px 0 2px 4em;line-height:1.6'
    _PARA_NUM_RE = _re_isi.compile(r'^\d+\.\s*')
    isi_html = ""
    _para_num = 2
    for para_text in [p.strip() for p in isi_raw.split('\n\n') if p.strip()]:
        # Skip duplicate "Dengan segala hormatnya" — already shown as fixed opening line
        if _DENGAN_RE.match(_strip_para_num(para_text)):
            continue
        lines = para_text.split('\n')
        is_first = True
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if is_first:
                line_s = _strip_para_num(line)
                if _TARIKH_BARIS_RE.match(line_s):
                    isi_html += f'<p style="{_indent_p}">{line_s}</p>'
                else:
                    # Tambah nombor jika belum ada (contoh: surat pemakluman dari PDF)
                    if _PARA_NUM_RE.match(line):
                        display = line  # sudah ada nombor, guna terus
                    else:
                        display = f'{_para_num}. {line_s}'
                    isi_html += f'<p style="{_plain_p}">{display}</p>'
                is_first = False
            else:
                line_s = _strip_para_num(line)
                if _TARIKH_BARIS_RE.match(line_s):
                    isi_html += f'<p style="{_indent_p}">{line_s}</p>'
                else:
                    isi_html += f'<p style="{_plain_p}">{line_s}</p>'
        isi_html += '<p style="margin:0 0 4px 0"></p>'
        _para_num += 1

    n = 'style="margin:6px 0;line-height:1.6"'
    # Ruj.Kami+Tarikh right-aligned, then address left — per template
    penerima_lines = [penerima_line, f.get("penerima_organisasi", ""), f.get("penerima_alamat", "")]
    penerima_html = "<br>".join(l for l in penerima_lines if l)
    rujukan_html = (
        f'Ruj. Kami &nbsp;: {f.get("rujukan","")}<br>'
        f'Tarikh &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;: {f.get("tarikh","")}'
    )

    sig_extra = ""
    if f.get("nama_pejabat"):
        sig_extra += f'<p style="margin:2px 0">{f.get("nama_pejabat","")}</p>'
    if f.get("nama_organisasi"):
        sig_extra += f'<p style="margin:2px 0">{f.get("nama_organisasi","")}</p>'

    return (
        f'<div style="font-family:Arial,sans-serif;font-size:12pt;line-height:1.5;color:#000">'
        f'{lh_block}'
        f'<hr style="border:none;border-top:2px solid #000;margin:8px 0 14px 0">'
        f'<p style="text-align:right;line-height:1.8;margin:0 0 10px 0">{rujukan_html}</p>'
        f'<p style="line-height:1.8;margin:0 0 14px 0">{penerima_html}</p>'
        f'<p {n}>{panggilan},</p>'
        f'<p style="margin:10px 0;font-weight:bold;text-align:left">{f.get("tajuk","").upper()}</p>'
        f'<p style="margin:8px 0;line-height:1.6;text-align:justify">Dengan segala hormatnya perkara di atas adalah dirujuk.</p>'
        f'{isi_html}'
        f'<div style="page-break-inside:avoid">'
        f'<p {n}>Sekian, terima kasih.</p>'
        f'<p style="margin:11pt 0 0 0;line-height:1.6"><b>&ldquo;MALAYSIA MADANI&rdquo;</b></p>'
        f'<p style="margin:11pt 0 0 0;line-height:1.6"><b>&ldquo;BERKHIDMAT UNTUK NEGARA&rdquo;</b></p>'
        f'<p style="margin:11pt 0 0 0;line-height:1.6">Saya yang menjalankan amanah,</p>'
        f'<p style="margin:33pt 0 0 0;line-height:1.6"><b>({f.get("penandatangan_nama","").upper()})</b></p>'
        f'<p style="margin:2px 0">{f.get("penandatangan_jawatan","")}</p>'
        f'{sig_extra}'
        f'</div>'
        f'{sk_html}'
        f'{_build_senarai_edaran_html(f) if is_se else ""}'
        f'</div>'
    )


def _build_senarai_edaran_html(f: dict) -> str:
    """Jana HTML halaman senarai edaran — ALL CAPS, BOLD."""
    se_raw = str(f.get('senarai_edaran', '') or '')
    if not se_raw or not se_raw.strip():
        return ""
    entries = []
    try:
        data = json.loads(se_raw)
        if isinstance(data, list):
            entries = data
    except Exception:
        for line in se_raw.strip().splitlines():
            line = line.strip()
            if line:
                entries.append({"nama": "", "jawatan": "", "jabatan": line})
    if not entries:
        return ""

    has_nama = any(str(e.get("nama", "") or "").strip() for e in entries)

    rows_html = ""
    for i, entry in enumerate(entries, 1):
        nama    = str(entry.get("nama", "") or "").upper()
        jawatan = str(entry.get("jawatan", "") or "").upper()
        jabatan = str(entry.get("jabatan", "") or "").upper()
        nama_td = f'<td style="padding:5px 8px;font-weight:bold">{nama}</td>' if has_nama else ""
        rows_html += (
            f'<tr style="border-bottom:1px solid #ccc">'
            f'<td style="padding:5px 8px;font-weight:bold;width:40px">{i}.</td>'
            f'{nama_td}'
            f'<td style="padding:5px 8px;font-weight:bold">{jawatan}</td>'
            f'<td style="padding:5px 8px;font-weight:bold">{jabatan}</td>'
            f'</tr>'
        )

    nama_th = '<th style="padding:5px 8px;text-align:left">NAMA</th>' if has_nama else ""
    return (
        f'<div style="margin-top:40px;border-top:2px solid #000;padding-top:16px;page-break-before:always">'
        f'<p style="font-size:13pt;font-weight:bold;text-align:center;letter-spacing:2px;margin:0 0 16px 0">SENARAI EDARAN</p>'
        f'<table style="width:100%;border-collapse:collapse;font-family:Arial,sans-serif;font-size:11pt">'
        f'<thead><tr style="border-bottom:2px solid #000;background:#f0f0f0">'
        f'<th style="padding:5px 8px;text-align:left;width:40px">BIL.</th>'
        f'{nama_th}'
        f'<th style="padding:5px 8px;text-align:left">JAWATAN</th>'
        f'<th style="padding:5px 8px;text-align:left">JABATAN / SEKOLAH</th>'
        f'</thead><tbody>{rows_html}</tbody></table>'
        f'</div>'
    )


def _build_memo_docx(doc, fields: dict):
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _p(text="", bold=False, indent_cm=0, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.alignment = align
        if indent_cm:
            para.paragraph_format.left_indent = Cm(indent_cm)
        if text:
            run = para.add_run(text)
            run.font.size = Pt(size)
            run.font.name = "Arial"
            run.bold = bold
        return para

    def _cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.alignment = align
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Arial"
        run.bold = bold

    def _no_border_table(tbl):
        tbl_element = tbl._tbl
        tbl_pr = tbl_element.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl_element.insert(0, tbl_pr)
        tbl_borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'none')
            tbl_borders.append(b)
        tbl_pr.append(tbl_borders)

    def _add_hr(para):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _shade_cell(cell, hex_color='000000'):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tc_pr.append(shd)

    # ── "MEMO DALAMAN" title in black box ──
    title_tbl = doc.add_table(rows=1, cols=1)
    _no_border_table(title_tbl)
    title_cell = title_tbl.cell(0, 0)
    _shade_cell(title_cell, '000000')
    _cell_text(title_cell, "MEMO DALAMAN", bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    title_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    title_cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    title_cell.paragraphs[0].paragraph_format.space_after = Pt(4)

    doc.add_paragraph("")

    # ── Header info table (no borders, 3 cols: label | colon+value) ──
    ahli_str = fields.get('ahli', '')
    ahli_list = [a.strip() for a in ahli_str.split(',') if a.strip()] if ahli_str else ['']
    pengerusi = fields.get('pengerusi', '')

    # Build rows: (label, value) pairs
    rows_data = []
    if pengerusi:
        rows_data.append(("Kepada", pengerusi))
        for ahli in ahli_list:
            if ahli:
                rows_data.append(("", ahli))
    rows_data.append(("Daripada", fields.get('urus_setia', '')))
    rows_data.append(("Tarikh", fields.get('tarikh', '')))
    rows_data.append(("Ruj. Kami", fields.get('rujukan', '')))
    rows_data.append(("Perkara", fields.get('tajuk', '').upper()))

    info_tbl = doc.add_table(rows=len(rows_data), cols=3, style='Table Grid')

    COL_W = [Cm(3.5), Cm(0.5), Cm(13)]
    for i, (label, value) in enumerate(rows_data):
        row = info_tbl.rows[i]
        row.cells[0].width = COL_W[0]
        row.cells[1].width = COL_W[1]
        row.cells[2].width = COL_W[2]

        is_perkara = label == "Perkara"
        label_bold = bool(label)

        _cell_text(row.cells[0], label, bold=label_bold, size=12)
        _cell_text(row.cells[1], ":" if label else "", size=12)
        _cell_text(row.cells[2], value, bold=is_perkara,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY if is_perkara else WD_ALIGN_PARAGRAPH.LEFT,
                   size=12)


    doc.add_paragraph("")

    # ── Body ──
    panggilan = _auto_panggilan(fields.get('pengerusi', 'Tuan'))
    p_panggilan = _p(f"{panggilan},")
    p_panggilan.paragraph_format.space_before = Pt(6)
    p_panggilan.paragraph_format.space_after = Pt(6)

    p_dengan = _p("Dengan segala hormatnya perkara di atas adalah dirujuk.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p_dengan.paragraph_format.space_after = Pt(6)

    _isi_raw = _strip_para_num(fields.get('isi', '').split('\n')[0].strip())
    # Strip tarikh/masa/tempat sentences already handled by dedicated block
    for _kw in [fields.get('tarikh_acara', ''), fields.get('masa_acara', ''), fields.get('tempat_acara', '')]:
        if _kw and _kw.lower() in _isi_raw.lower():
            _idx = _isi_raw.lower().find(_kw.lower())
            _start = _isi_raw.rfind('.', 0, _idx)
            if _start != -1:
                _isi_raw = _isi_raw[:_start + 1].strip()
    isi = _isi_raw
    isi_para = doc.add_paragraph()
    isi_para.paragraph_format.space_after = Pt(6)
    isi_para.paragraph_format.space_before = Pt(0)
    isi_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r2 = isi_para.add_run(f"2.\t{isi}")
    r2.font.size = Pt(12); r2.font.name = "Arial"

    # Acara block
    for label, key in [("Tarikh", "tarikh_acara"), ("Masa", "masa_acara"), ("Tempat", "tempat_acara")]:
        val = fields.get(key, '')
        acara_p = _p(f"{label:<8}: {val}", indent_cm=2)
        acara_p.paragraph_format.space_after = Pt(2)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    p3.paragraph_format.space_before = Pt(6)
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.left_indent = Cm(1.27)
    p3.paragraph_format.first_line_indent = Cm(-1.27)
    r3 = p3.add_run("3.\tKehadiran tuan/puan pada tarikh dan masa yang ditetapkan amatlah dihargai.")
    r3.font.size = Pt(12); r3.font.name = "Arial"

    langkah_str = fields.get('langkah_kerja', '')
    langkah_list = [l.strip() for l in langkah_str.split('\n') if l.strip()] if langkah_str else []
    for i, step in enumerate(langkah_list):
        num = i + 4
        lk_para = doc.add_paragraph()
        lk_para.paragraph_format.space_after = Pt(6)
        lk_para.paragraph_format.space_before = Pt(0)
        lk_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lk_para.paragraph_format.left_indent = Cm(1.27)
        lk_para.paragraph_format.first_line_indent = Cm(-1.27)
        lk_run = lk_para.add_run(f"{num}.\t{step}")
        lk_run.font.size = Pt(12)
        lk_run.font.name = "Arial"

    doc.add_paragraph("")
    _p("Sekian, terima kasih.")
    doc.add_paragraph("")
    _p('"MALAYSIA MADANI"', bold=True)
    doc.add_paragraph("")
    _p('"BERKHIDMAT UNTUK NEGARA"', bold=True)
    doc.add_paragraph("")
    _p("Saya yang menjalankan amanah,")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    _p(f"({fields.get('penandatangan_nama', '').upper()})", bold=True)
    _p(fields.get('penandatangan_jawatan', ''))
    _p(fields.get('nama_pejabat', ''))


def get_session_info(session_id: str) -> dict | None:
    s = _get_session(session_id)
    return s if s.get("doc_type") else None


def send_email(session_id: str, to_email: str, subject: str) -> dict:
    document = get_document(session_id)
    if not document:
        return {"ok": False, "error": "Tiada dokumen yang sedia untuk dihantar. Sila lengkapkan dokumen dahulu."}

    sender_email = os.getenv("SENDER_EMAIL")
    sender_password = os.getenv("SENDER_APP_PASSWORD")

    if not sender_email or not sender_password:
        return {"ok": False, "error": "Konfigurasi emel belum ditetapkan. Sila tetapkan SENDER_EMAIL dan SENDER_APP_PASSWORD dalam .env."}

    try:
        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = to_email
        msg["Subject"] = subject

        msg.attach(MIMEText(document, "plain", "utf-8"))

        session = _get_session(session_id)
        doc_type = session.get("doc_type", "dokumen")
        filename = f"{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        attachment = MIMEBase("application", "octet-stream")
        attachment.set_payload(document.encode("utf-8"))
        encoders.encode_base64(attachment)
        attachment.add_header("Content-Disposition", f"attachment; filename={filename}")
        msg.attach(attachment)

        smtp_host = os.getenv("SMTP_HOST", "smtp.gmail.com")
        smtp_port = int(os.getenv("SMTP_PORT", "587"))

        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)

        return {"ok": True, "message": f"Emel berjaya dihantar ke {to_email}."}
    except Exception as e:
        return {"ok": False, "error": f"Gagal menghantar emel: {e}"}


def clear_session(session_id: str):
    from backend.session_store import get_store
    get_store().delete_ns(session_id, _NS)


def _try_parse_json(text: str) -> dict | None:
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        lines = lines[1:] if lines[0].startswith("```") else lines
        end = next((i for i, l in enumerate(lines) if l.strip() == "```"), len(lines))
        text = "\n".join(lines[:end])
    try:
        return json.loads(text)
    except (json.JSONDecodeError, ValueError):
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end > start:
            try:
                return json.loads(text[start:end + 1])
            except (json.JSONDecodeError, ValueError):
                pass
    return None
