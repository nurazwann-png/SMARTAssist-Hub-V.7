import json
import os
import uvicorn
from dotenv import load_dotenv

load_dotenv()


def _parse_agent_json(text: str):
    """Parse JSON from agent output, stripping markdown code fences if present."""
    t = text.strip()
    if t.startswith("```"):
        lines = t.split("\n")
        lines = lines[1:] if lines[0].startswith("```") else lines
        end = next((i for i, l in enumerate(lines) if l.strip() == "```"), len(lines))
        t = "\n".join(lines[:end])
    try:
        return json.loads(t)
    except (json.JSONDecodeError, ValueError):
        s, e = text.find("{"), text.rfind("}")
        if s != -1 and e > s:
            try:
                return json.loads(text[s:e + 1])
            except (json.JSONDecodeError, ValueError):
                pass
    return None


def _shorten_doc_message(structured, lang="ms"):
    """Replace the long dumped document in `message` with a short notice.

    The LLM sometimes embeds the whole document in the `message` field, which
    then renders in full above the preview. Since the full document already
    appears in the pratonton (document_preview / corrected_document), keep the
    message area to a short notification. Applies to all three document agents.
    """
    if not structured:
        return
    en = (lang == "en")
    if structured.get("document_preview"):
        structured["message"] = (
            "✅ Document ready. Please review the preview below."
            if en else
            "✅ Dokumen telah dijana. Sila semak pratonton di bawah."
        )
    elif structured.get("corrected_document"):
        structured["message"] = (
            "✅ Document corrected. Please review the preview below."
            if en else
            "✅ Dokumen telah diperbetulkan. Sila semak pratonton di bawah."
        )
from datetime import datetime
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from jinja2 import Environment, FileSystemLoader
from pydantic import BaseModel
from starlette.middleware.sessions import SessionMiddleware

from fastapi.responses import PlainTextResponse, Response

from auth import router as auth_router, get_current_user
from backend.profile_store import get_profile, save_profile
from backend.orchestrator import run_query
from backend.session_store import get_store as _get_store
from agents.data_analysis import upload_file as da_upload, get_session_data as da_get_data
from agents.letter_generator import (
    get_document as lg_get_document,
    build_docx as lg_build_docx,
    send_email as lg_send_email,
    get_session_info as lg_get_session,
    inject_pdf_context as lg_inject_pdf_context,
)
from agents.report_generator import (
    get_document as rg_get_document,
    build_docx as rg_build_docx,
    send_email as rg_send_email,
    get_session_info as rg_get_session,
    add_report_image as rg_add_image,
    remove_report_image as rg_remove_image,
    get_report_images as rg_get_images,
)

app = FastAPI(title="SMARTAssist Hub", version="7.0")

# Session middleware must be added before including routers
app.add_middleware(
    SessionMiddleware,
    secret_key=os.getenv("SESSION_SECRET_KEY", "change-me-in-production"),
)

app.include_router(auth_router)
app.mount("/static", StaticFiles(directory="static"), name="static")

_jinja_env = Environment(loader=FileSystemLoader("templates"), autoescape=True)

_feedback: list[dict] = []

AGENT_LABELS = {
    "data_analysis": {"icon": "\U0001f4ca", "name": "Analisis Data"},
    "letter_generator": {"icon": "✉️", "name": "Penjana Surat"},
    "report_generator": {"icon": "\U0001f4cb", "name": "Penjana Laporan"},
    "document_reviewer": {"icon": "\U0001f4dd", "name": "Semakan Dokumen"},
    "kpm_support": {"icon": "\U0001f6df", "name": "Sokongan KPM"},
    "fallback": {"icon": "\U0001f4ac", "name": "Pembantu Umum"},
}


class ChatRequest(BaseModel):
    message: str
    session_id: str = "default"


class FeedbackRequest(BaseModel):
    session_id: str
    message_index: int
    feedback: str  # 'up' or 'down'
    agent: str = ""


@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request, error: str = "", email: str = ""):
    user = get_current_user(request)
    if user:
        return RedirectResponse("/")
    template = _jinja_env.get_template("login.html")
    html = template.render(error=error, email=email)
    return HTMLResponse(html)


@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    user = get_current_user(request)
    if not user:
        return RedirectResponse("/login")
    profile = get_profile(user["sub"])
    agents_list = [
        {"key": k, "icon": v["icon"], "name": v["name"]}
        for k, v in AGENT_LABELS.items()
        if k != "fallback"
    ]
    template = _jinja_env.get_template("index.html")
    html = template.render(agents_list=agents_list, user=user, profile=profile)
    return HTMLResponse(html)


class AgentChatRequest(BaseModel):
    message: str
    session_id: str = "default"
    agent: str = ""
    lang: str = "bm"


@app.post("/api/chat")
async def chat(req: ChatRequest):
    store = _get_store()
    history = store.get_messages(req.session_id)
    store.append_message(req.session_id, {"role": "user", "content": req.message})
    history.append({"role": "user", "content": req.message})

    result = run_query(query=req.message, conversation_history=history, session_id=req.session_id)

    agent = result.get("active_agent", "fallback")
    output = result.get("agent_output", "Tiada respons.")

    store.append_message(req.session_id, {"role": "assistant", "content": output, "agent": agent})
    meta = store.get_meta(req.session_id) or {}
    store.upsert_meta(
        req.session_id,
        agent=agent,
        title=meta.get("title") or req.message[:60],
    )

    agent_info = AGENT_LABELS.get(agent, {"icon": "\U0001f4ac", "name": agent})

    structured = None
    if agent in ("data_analysis", "letter_generator", "report_generator", "document_reviewer"):
        structured = _parse_agent_json(output)
        _shorten_doc_message(structured, getattr(req, "lang", "ms"))

    return JSONResponse({
        "response": output,
        "agent": agent,
        "agent_icon": agent_info["icon"],
        "agent_name": agent_info["name"],
        "structured": structured,
    })


@app.post("/api/upload")
async def upload(file: UploadFile = File(...), session_id: str = Form("default")):
    contents = await file.read()
    result = da_upload(contents, file.filename, session_id)

    if result["ok"]:
        _get_store().append_message(session_id, {
            "role": "assistant",
            "content": f"Fail '{result['filename']}' berjaya dimuat naik. "
                       f"{result['rows']} baris x {result['columns']} lajur. "
                       f"Lajur: {', '.join(result['column_names'][:15])}"
                       f"{'...' if len(result['column_names']) > 15 else ''}",
            "agent": "data_analysis",
        })

    return JSONResponse(result)


def _text_to_review_html(text: str) -> str:
    """Fallback: wrap extracted plain text into simple paragraph HTML when the
    layout-aware PDF reconstruction is unavailable. Each line becomes a <p> so
    inline issue-badges map per line."""
    import html as _html
    parts = ['<div style="font-family:Arial,sans-serif;font-size:12pt;line-height:1.6;color:#000">']
    for ln in text.split("\n"):
        s = ln.strip()
        if not s:
            parts.append('<p style="margin:6px 0">&nbsp;</p>')
        else:
            parts.append(f'<p style="margin:6px 0">{_html.escape(s)}</p>')
    parts.append('</div>')
    return "".join(parts)


def _pdf_page_lines(page):
    """Reconstruct a PDF page's lines with alignment/bold/spacing from word
    positions, so the reviewer preview keeps the uploaded letter's format
    (right/centre alignment, bold headings, paragraph spacing)."""
    import statistics
    words = page.extract_words(use_text_flow=False, keep_blank_chars=False,
                               extra_attrs=["fontname"])
    if not words:
        return []
    groups = {}
    for w in words:
        groups.setdefault(round(w["top"]), []).append(w)
    body_left = min(w["x0"] for w in words)
    right_edge = max(w["x1"] for w in words)
    tops = sorted(groups)
    deltas = [b - a for a, b in zip(tops, tops[1:])]
    pitch = statistics.median(deltas) if deltas else 16.0
    out = []
    prev_top = None
    for top in tops:
        ws = sorted(groups[top], key=lambda w: w["x0"])
        x0 = min(w["x0"] for w in ws); x1 = max(w["x1"] for w in ws)
        left_gap = x0 - body_left; right_gap = right_edge - x1
        align = "left" if left_gap < 15 else ("right" if right_gap < 15 else "center")
        bold = sum(1 for w in ws if "Bold" in (w.get("fontname") or "")) >= max(1, len(ws) // 2)
        blanks = 0
        if prev_top is not None and pitch:
            blanks = min(2, max(0, round((top - prev_top) / pitch) - 1))
        prev_top = top
        out.append({
            "text": " ".join(w["text"] for w in ws),
            "align": align, "bold": bold, "blanks": blanks,
        })
    return out


# In-memory cache of the most recent uploaded PDF bytes per session, so the
# review step can locate each issue's text and highlight it on the page images.
_REVIEW_PDF_CACHE = {}


def _locate_issue_boxes(raw_pdf, issues):
    """Attach a normalised highlight box {page,x,y,w,h} to each issue whose
    verbatim `petikan` text can be found in the uploaded PDF."""
    import io as _io3
    import pdfplumber
    try:
        pdf = pdfplumber.open(_io3.BytesIO(raw_pdf))
    except Exception:
        return
    try:
        pages = pdf.pages
        for issue in issues:
            petikan = (issue.get("petikan") or "").strip()
            if not petikan:
                continue
            words = petikan.split()
            attempts = [petikan] + [" ".join(words[:n]) for n in (6, 4, 3) if n < len(words)]
            found = None
            for attempt in attempts:
                for pi, page in enumerate(pages):
                    try:
                        res = page.search(attempt, case=False)
                    except Exception:
                        res = []
                    if res:
                        r = res[0]
                        W = page.width or 1.0
                        H = page.height or 1.0
                        found = {
                            "page": pi,
                            "x": round(r["x0"] / W, 4),
                            "y": round(r["top"] / H, 4),
                            "w": round((r["x1"] - r["x0"]) / W, 4),
                            "h": round((r["bottom"] - r["top"]) / H, 4),
                        }
                        break
                if found:
                    break
            if found:
                issue["highlight"] = found
    finally:
        pdf.close()


def _pdf_to_page_images(pdf, resolution=150):
    """Render each PDF page to a PNG data-URI. This gives a preview that is a
    pixel-faithful image of the uploaded PDF and displays reliably in any browser
    (plain <img>), avoiding the blob/iframe PDF-viewer that renders blank."""
    import io as _io2, base64 as _b64
    out = []
    for page in pdf.pages:
        im = page.to_image(resolution=resolution)
        buf = _io2.BytesIO()
        im.save(buf, format="PNG")
        out.append("data:image/png;base64," + _b64.b64encode(buf.getvalue()).decode("ascii"))
    return out


def _pdf_to_review_html(pdf) -> str:
    """Build editable review HTML from an open pdfplumber PDF, following the
    uploaded file's own layout as closely as possible: per-line alignment
    (left/right/centre), bold and paragraph spacing are taken directly from the
    PDF's word positions, without any KPM-format normalisation."""
    import html as _html
    parts = ['<div style="font-family:Arial,sans-serif;font-size:12pt;line-height:1.5;color:#000">']
    any_line = False
    for page in pdf.pages:
        for ln in _pdf_page_lines(page):
            any_line = True
            for _ in range(ln["blanks"]):
                parts.append('<p style="margin:0;line-height:1.4">&nbsp;</p>')
            body = _html.escape(ln["text"])
            if ln["bold"]:
                body = f"<b>{body}</b>"
            parts.append(f'<p style="margin:2px 0;text-align:{ln["align"]}">{body}</p>')
    parts.append("</div>")
    return "".join(parts) if any_line else ""


@app.post("/api/review/upload")
async def review_upload(file: UploadFile = File(...), session_id: str = Form("default")):
    """Extract text from PDF or DOCX and store in document reviewer session."""
    import io as _io
    fname = file.filename or ""
    ext = _Path(fname).suffix.lower()
    allowed = {".pdf", ".docx", ".doc"}
    if ext not in allowed:
        return JSONResponse({"ok": False, "error": f"Format tidak disokong: {ext}. Sila gunakan PDF atau Word (.docx)."}, status_code=400)

    raw = await file.read()
    text = ""
    doc_html = None   # rich HTML for visual preview
    doc_type = "Dokumen"
    pdf_images = []   # per-page PNG data-URIs for faithful PDF preview
    _REVIEW_PDF_CACHE.pop(session_id, None)   # clear any previous upload

    try:
        if ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(_io.BytesIO(raw)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
                # Render each page to an image so the preview is pixel-faithful to
                # the uploaded PDF and displays reliably (plain <img>, not an
                # iframe/blob PDF viewer which renders blank in-app).
                try:
                    pdf_images = _pdf_to_page_images(pdf)
                except Exception:
                    pdf_images = []
            text = "\n".join(p for p in pages if p.strip())
            doc_type = "PDF"
            doc_html = None
            # Cache raw bytes so the review step can locate & highlight issues
            _REVIEW_PDF_CACHE[session_id] = raw
        elif ext in (".docx", ".doc"):
            # mammoth for rich HTML preview (preserves tables, bold, etc.)
            import mammoth
            result = mammoth.convert_to_html(_io.BytesIO(raw))
            doc_html = result.value  # HTML string

            # python-docx for plain text (AI review)
            from docx import Document as _DocxDoc
            doc = _DocxDoc(_io.BytesIO(raw))
            parts = []
            for para in doc.paragraphs:
                parts.append(para.text)
            # Also extract table cells so AI sees tabular content
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        parts.append("  ".join(row_texts))
            text = "\n".join(parts)
            doc_type = "Word (.docx)"
    except Exception as e:
        return JSONResponse({"ok": False, "error": f"Gagal membaca fail: {e}"}, status_code=500)

    if not text.strip():
        return JSONResponse({"ok": False, "error": "Fail kosong atau teks tidak dapat diekstrak."}, status_code=400)

    # Detect actual document type from content (overrides file-format label).
    # Collapse whitespace around ':' and '/' so markers match regardless of
    # spacing (e.g. both "KEPADA:" and "KEPADA :").
    import re as _re_det
    _upper = _re_det.sub(r'\s*([:/])\s*', r'\1', text.upper())
    _report_markers = ["NAMA PROGRAM", "BUTIRAN PERLAKSANAAN", "ONE PAGE REPORT",
                       "DISEDIAKAN OLEH", "CADANGAN/TINDAKAN", "RUMUSAN", "TARIKH PROGRAM"]
    _letter_markers = ["RUJ. KAMI", "MALAYSIA MADANI", "SAYA YANG MENJALANKAN AMANAH",
                       "Y.BHG", "TUAN/PUAN", "BERKHIDMAT UNTUK NEGARA"]
    _memo_markers = ["MEMO DALAMAN", "MEMO RASMI", "KEPADA:", "DARIPADA:", "TARIKH:", "PERKARA:"]
    _report_score = sum(1 for m in _report_markers if m in _upper)
    _letter_score = sum(1 for m in _letter_markers if m in _upper)
    _memo_score   = sum(1 for m in _memo_markers if m in _upper)
    if _report_score >= 2 and _report_score >= _letter_score and _report_score >= _memo_score:
        doc_type = "Laporan Satu Muka Surat"
    elif _memo_score >= 2 and _memo_score >= _letter_score:
        doc_type = "Memo Dalaman"
    elif _letter_score >= 2:
        doc_type = "Surat Rasmi"
    # else keep original doc_type (PDF / Word)

    from agents.document_reviewer import set_uploaded_document
    set_uploaded_document(session_id, text, fname, doc_type, html=doc_html)

    # Add to session history so context is preserved
    _get_store().append_message(session_id, {"role": "user", "content": f"[Fail dimuat naik: {fname}]"})

    return JSONResponse({
        "ok": True,
        "filename": fname,
        "doc_type": doc_type,
        "char_count": len(text),
        "preview": text[:300],
        "text": text,
        "html": doc_html,   # rich HTML (DOCX); None for PDF (shown as page images)
        "is_pdf": ext == ".pdf" and doc_html is None,
        "pdf_images": pdf_images,   # per-page PNG data-URIs for faithful PDF preview
    })


@app.post("/api/letter/upload-pdf")
async def letter_upload_pdf(file: UploadFile = File(...), session_id: str = Form("default")):
    """Ekstrak teks dari PDF dan pra-isi session letter_generator untuk jana surat iringan."""
    import io as _io
    fname = file.filename or ""
    ext = _Path(fname).suffix.lower()

    if ext != ".pdf":
        return JSONResponse({"ok": False, "error": "Hanya fail PDF disokong. Sila muat naik fail .pdf."}, status_code=400)

    raw = await file.read()

    # Had saiz — 10MB
    if len(raw) > 10 * 1024 * 1024:
        return JSONResponse({"ok": False, "error": "Fail terlalu besar (melebihi 10MB). Sila muat naik PDF yang lebih kecil."}, status_code=400)

    # Ekstrak teks dalam memori — tidak simpan ke disk
    try:
        import pdfplumber
        with pdfplumber.open(_io.BytesIO(raw)) as pdf:
            # Ambil maksimum 5 halaman pertama sahaja
            pages_text = [p.extract_text() or "" for p in pdf.pages[:5]]
            total_pages = len(pdf.pages)
    except Exception as e:
        err_lower = str(e).lower()
        if "password" in err_lower or "encrypt" in err_lower:
            return JSONResponse({"ok": False, "error": "PDF ini dilindungi kata laluan dan tidak dapat diproses."}, status_code=400)
        return JSONResponse({"ok": False, "error": f"Gagal membaca PDF: {e}"}, status_code=500)

    full_text = "\n".join(p for p in pages_text if p.strip())

    if not full_text.strip():
        return JSONResponse({
            "ok": False,
            "error": "PDF ini adalah PDF imbasan/imej. Teks tidak dapat diekstrak. Sila gunakan PDF yang mempunyai teks boleh pilih."
        }, status_code=400)

    # Had teks untuk LLM — ambil 4000 aksara pertama
    text_for_llm = full_text[:4000]
    pages_note = f" (hanya 5 halaman pertama daripada {total_pages} halaman dianalisis)" if total_pages > 5 else ""

    # LLM analisis — ekstrak field dan jana isi surat pemakluman
    from backend.deepseek_client import chat_completion
    analysis_prompt = f"""Analisis teks dari dokumen rasmi berikut dan ekstrak maklumat dalam format JSON.

Dokumen yang dimuat naik ini akan digunakan untuk menjana SURAT PEMAKLUMAN — iaitu surat rasmi untuk memaklumkan penerima baharu tentang kandungan dokumen asal ini.

Kembalikan HANYA JSON ini tanpa sebarang teks lain:
{{
  "suggested_type": "surat atau memo",
  "rujukan": "nombor rujukan rasmi dokumen asal jika ada (cth: KPM.600-1/2/3), atau null",
  "tarikh": "tarikh dalam format D Bulan YYYY jika ada (cth: 10 Julai 2026), atau null",
  "tajuk": "PEMAKLUMAN: [tajuk/perkara utama dokumen asal]",
  "penerima_nama": "nama penerima jika ada (untuk surat), atau null",
  "penerima_organisasi": "nama organisasi/jabatan penerima jika ada, atau null",
  "isi_pemakluman": "perenggan isi surat pemakluman dalam Bahasa Melayu rasmi (2-3 ayat): mula dengan 'Adalah dimaklumkan bahawa...', kemudian ringkaskan kandungan utama dokumen asal, dan nyatakan tindakan atau makluman yang diperlukan daripada penerima. Gunakan gaya bahasa surat rasmi KPM.",
  "analysis_summary": "ringkasan pendek 1-2 ayat dalam Bahasa Melayu tentang kandungan dokumen ini"
}}

Teks dokumen:
{text_for_llm}"""

    extracted = {}
    analysis_summary = ""
    suggested_type = "surat"

    try:
        raw_json = chat_completion(
            messages=[{"role": "user", "content": analysis_prompt}],
            temperature=0.1,
            max_tokens=800
        )
        import re as _re
        # Ekstrak JSON dari response (buang teks sekeliling jika ada)
        json_match = _re.search(r'\{[\s\S]*\}', raw_json)
        if json_match:
            parsed = json.loads(json_match.group())
            suggested_type = parsed.get("suggested_type", "surat")
            if suggested_type not in ("surat", "memo"):
                suggested_type = "surat"
            analysis_summary = parsed.get("analysis_summary", "")
            for k in ("rujukan", "tarikh", "tajuk", "penerima_nama", "penerima_organisasi", "isi_pemakluman"):
                val = parsed.get(k)
                if val and str(val).strip().lower() not in ("null", "none", ""):
                    extracted[k] = str(val).strip()
    except Exception:
        pass  # Fallback: teruskan tanpa field pra-isi

    # Inject ke session letter_generator
    lg_inject_pdf_context(session_id, extracted, suggested_type)

    return JSONResponse({
        "ok": True,
        "filename": fname,
        "suggested_type": suggested_type,
        "extracted_fields": extracted,
        "analysis_summary": analysis_summary,
        "char_count": len(full_text),
        "pages_note": pages_note,
    })


def _render_html_to_pdf(html_content: str) -> bytes:
    """Render the document-preview HTML into a PDF using reportlab.

    Supports the subset of HTML the app produces: <img> (letterhead), <hr>,
    <p style="..."> with text-align / margin / padding-left / text-indent /
    font-weight / line-height, inline <b>/<i>/<u>/<br>, headings, lists and
    tables. Mirrors the on-screen preview so the PDF matches the .docx output.
    """
    import io as _io
    import re as _re
    import base64 as _b64
    from bs4 import BeautifulSoup, NavigableString
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
        Table, TableStyle, HRFlowable,
    )
    from reportlab.lib.styles import ParagraphStyle

    FONT_PT = 12.0
    MARGIN = 25.4 * mm
    CONTENT_W = A4[0] - 2 * MARGIN

    def _css_len(val, font_pt=FONT_PT):
        if not val:
            return 0.0
        m = _re.match(r'^\s*(-?[\d.]+)\s*(em|rem|px|pt)?\s*$', val)
        if not m:
            return 0.0
        num = float(m.group(1)); unit = m.group(2) or 'px'
        if unit in ('em', 'rem'):
            return num * font_pt
        if unit == 'px':
            return num * 0.75
        return num

    def _parse_style(s):
        d = {}
        for part in (s or '').split(';'):
            if ':' in part:
                k, v = part.split(':', 1)
                d[k.strip().lower()] = v.strip()
        return d

    def _margin_tb(val):
        toks = (val or '').split()
        v = [_css_len(t) for t in toks]
        if not v:
            return 0.0, 0.0
        if len(v) == 1:
            return v[0], v[0]
        if len(v) == 2:
            return v[0], v[0]
        if len(v) == 3:
            return v[0], v[2]
        return v[0], v[2]

    def _esc(t):
        return t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    def _inline(node):
        out = []
        for c in node.children:
            if isinstance(c, NavigableString):
                out.append(_esc(str(c)))
            else:
                nm = (c.name or '').lower()
                inner = _inline(c)
                if nm in ('b', 'strong'):
                    out.append(f'<b>{inner}</b>')
                elif nm in ('i', 'em'):
                    out.append(f'<i>{inner}</i>')
                elif nm == 'u':
                    out.append(f'<u>{inner}</u>')
                elif nm == 'br':
                    out.append('<br/>')
                else:
                    out.append(inner)
        return ''.join(out)

    def _margin_left(val):
        """Extract left margin from CSS margin shorthand (top right bottom left)."""
        toks = (val or '').split()
        if len(toks) == 4:
            return _css_len(toks[3])
        if len(toks) == 2:
            return _css_len(toks[1])
        return 0.0

    def _para_style(st, default_size=FONT_PT, bold_default=False):
        align = {'right': TA_RIGHT, 'center': TA_CENTER,
                 'justify': TA_JUSTIFY, 'left': TA_LEFT}.get(
            st.get('text-align', 'left'), TA_LEFT)
        lh = st.get('line-height', '')
        leading = (float(lh) * default_size) if _re.match(r'^[\d.]+$', lh or '') else default_size * 1.4
        bold = bold_default or st.get('font-weight', '') in ('bold', '700')
        mt, mb = _margin_tb(st.get('margin', ''))
        # Left indent: explicit padding-left takes priority, else margin-left / margin shorthand
        pl = st.get('padding-left', '')
        ml_explicit = st.get('margin-left', '')
        left_indent = (
            _css_len(pl) if pl else
            _css_len(ml_explicit) if ml_explicit else
            _margin_left(st.get('margin', ''))
        )
        return ParagraphStyle(
            'p',
            fontName='Helvetica-Bold' if bold else 'Helvetica',
            fontSize=default_size,
            leading=leading,
            alignment=align,
            leftIndent=left_indent,
            firstLineIndent=_css_len(st.get('text-indent', '0')),
            spaceBefore=mt,
            spaceAfter=mb,
        )

    def _css_to_pt(val):
        """Convert a CSS length string (cm/mm/pt/px/in) to reportlab points."""
        if not val:
            return None
        m = _re.match(r'([\d.]+)\s*(cm|mm|pt|px|in)', val.strip())
        if not m:
            return None
        n, u = float(m.group(1)), m.group(2)
        if u == 'cm': return n * 28.35
        if u == 'mm': return n * 2.835
        if u == 'pt': return n
        if u == 'px': return n * 0.75
        if u == 'in': return n * 72
        return None

    def _make_image(img):
        src = img.get('src', '') or ''
        raw = None
        try:
            if src.startswith('data:'):
                b64 = src.split(',', 1)[1]
                raw = _b64.b64decode(b64)
            elif '/api/letterhead/image/' in src:
                from backend.letterhead_store import _LH_DIR
                fn = src.rsplit('/', 1)[-1]
                p = _LH_DIR / fn
                if p.exists():
                    raw = p.read_bytes()
            if not raw:
                return None
            from PIL import Image as _PIL
            bio = _io.BytesIO(raw)
            im = _PIL.open(bio)
            iw, ih = im.size
            # Normalise to PNG so reportlab always reads it
            png = _io.BytesIO()
            im.convert('RGBA').save(png, format='PNG')
            png.seek(0)
            # Respect explicit CSS width/height on the <img> (e.g. width:6.71cm)
            st = _parse_style(img.get('style', ''))
            fixed_w = _css_to_pt(st.get('width', ''))
            fixed_h = _css_to_pt(st.get('height', ''))
            if fixed_w and fixed_h:
                w, h = fixed_w, fixed_h
            else:
                max_h = 150 * 0.75  # 150px cap fallback
                ratio = iw / ih if ih else 1
                h = min(max_h, ih * 0.75)
                w = h * ratio
                if w > CONTENT_W:
                    w = CONTENT_W
                    h = w / ratio
            rl = RLImage(png, width=w, height=h)
            rl.hAlign = 'CENTER'
            return rl
        except Exception:
            return None

    def _parse_color(val):
        """Convert CSS color (#rgb, #rrggbb, named) to reportlab Color."""
        if not val:
            return None
        val = val.strip().lower()
        named = {'black': colors.black, 'white': colors.white, 'red': colors.red,
                 'blue': colors.blue, 'green': colors.green, 'gray': colors.gray,
                 'grey': colors.gray, 'silver': colors.silver}
        if val in named:
            return named[val]
        if val.startswith('#'):
            h = val[1:]
            if len(h) == 3:
                h = h[0]*2 + h[1]*2 + h[2]*2
            if len(h) == 6:
                try:
                    r, g, b = int(h[0:2],16)/255, int(h[2:4],16)/255, int(h[4:6],16)/255
                    return colors.Color(r, g, b)
                except Exception:
                    pass
        return None

    def _make_table(tbl):
        # First pass: determine total column count and collect explicit widths
        num_cols = 1
        col_widths_px = {}  # col_index -> width in px
        for tr in tbl.find_all('tr'):
            cells = tr.find_all(['td', 'th'])
            total = sum(int(c.get('colspan', 1)) for c in cells)
            if total > num_cols:
                num_cols = total
            # Collect widths from first data row (non-spanning cells)
            ci = 0
            for c in cells:
                cs = int(c.get('colspan', 1))
                if cs == 1 and ci not in col_widths_px:
                    st = _parse_style(c.get('style', ''))
                    w_str = st.get('width', '') or c.get('width', '')
                    if w_str:
                        m = _re.match(r'([\d.]+)\s*(px)?', w_str)
                        if m:
                            col_widths_px[ci] = float(m.group(1))
                ci += cs

        # Build column widths: use explicit px widths where available
        if col_widths_px and len(col_widths_px) == num_cols:
            total_px = sum(col_widths_px.values())
            col_widths = [CONTENT_W * col_widths_px[i] / total_px for i in range(num_cols)]
        elif col_widths_px:
            # Partial: fixed px cols, rest shares remaining space
            known_pt = sum(_css_len(f'{col_widths_px[i]}px') for i in col_widths_px)
            remaining = max(CONTENT_W - known_pt, 10)
            unknown_cols = num_cols - len(col_widths_px)
            col_widths = []
            for i in range(num_cols):
                if i in col_widths_px:
                    col_widths.append(_css_len(f'{col_widths_px[i]}px'))
                else:
                    col_widths.append(remaining / max(unknown_cols, 1))
        else:
            col_widths = [CONTENT_W / num_cols] * num_cols

        rows = []
        span_cmds = []
        color_cmds = []
        row_idx = 0
        for tr in tbl.find_all('tr'):
            cells = tr.find_all(['td', 'th'])
            if not cells:
                continue
            row_data = []
            col_idx = 0
            for c in cells:
                cs = int(c.get('colspan', 1))
                rs = int(c.get('rowspan', 1))
                st = _parse_style(c.get('style', ''))
                bold_default = c.name == 'th' or st.get('font-weight', '') in ('bold', '700')
                # Apply text color from style
                txt_color = _parse_color(st.get('color', ''))
                cell_style = _para_style(st, 10.0, bold_default=bold_default)
                if txt_color:
                    cell_style = ParagraphStyle('pc', parent=cell_style, textColor=txt_color)
                para = Paragraph(_inline(c) or '&nbsp;', cell_style)
                row_data.append(para)
                if cs > 1 or rs > 1:
                    span_cmds.append(('SPAN', (col_idx, row_idx),
                                      (col_idx + cs - 1, row_idx + rs - 1)))
                # Background color
                bg = _parse_color(st.get('background', '') or st.get('background-color', ''))
                if bg:
                    color_cmds.append(('BACKGROUND', (col_idx, row_idx),
                                       (col_idx + cs - 1, row_idx + rs - 1), bg))
                # Text color for entire cell span
                if txt_color:
                    color_cmds.append(('TEXTCOLOR', (col_idx, row_idx),
                                       (col_idx + cs - 1, row_idx + rs - 1), txt_color))
                # Padding overrides from style
                for _ in range(cs - 1):
                    row_data.append(Paragraph('', _para_style({})))
                    col_idx += 1
                col_idx += 1
            while len(row_data) < num_cols:
                row_data.append(Paragraph('', _para_style({})))
            rows.append(row_data)
            row_idx += 1

        if not rows:
            return None

        t = Table(rows, hAlign='LEFT', colWidths=col_widths)
        style_cmds = [
            ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ] + span_cmds + color_cmds
        t.setStyle(TableStyle(style_cmds))
        return t

    flow = []
    from reportlab.platypus import PageBreak as _PB, KeepTogether as _KT

    def _walk(node, target=None):
        if target is None:
            target = flow
        for c in node.children:
            if isinstance(c, NavigableString):
                continue
            nm = (c.name or '').lower()
            if nm == 'img':
                el = _make_image(c)
                if el:
                    target.append(el)
            elif nm == 'hr':
                target.append(HRFlowable(width='100%', thickness=1.5,
                                         color=colors.black, spaceBefore=4, spaceAfter=8))
            elif nm == 'br':
                target.append(Spacer(1, FONT_PT * 1.6))
            elif nm == 'p':
                txt = _inline(c).strip()
                st = _parse_style(c.get('style', ''))
                if not txt:
                    mt, mb = _margin_tb(st.get('margin', ''))
                    target.append(Spacer(1, max(6.0, mt + mb)))
                else:
                    target.append(Paragraph(txt, _para_style(st)))
            elif nm in ('h1', 'h2', 'h3', 'h4'):
                size = {'h1': 18, 'h2': 15, 'h3': 13, 'h4': 12}[nm]
                st = _parse_style(c.get('style', ''))
                target.append(Paragraph(_inline(c) or '&nbsp;',
                                        _para_style(st, size, bold_default=True)))
            elif nm in ('ul', 'ol'):
                ordered = nm == 'ol'
                for i, li in enumerate(c.find_all('li', recursive=False), 1):
                    bullet = f'{i}.' if ordered else '•'
                    ps = ParagraphStyle('li', fontName='Helvetica', fontSize=FONT_PT,
                                        leading=FONT_PT * 1.4, leftIndent=18,
                                        firstLineIndent=-12, spaceAfter=2)
                    target.append(Paragraph(f'{bullet}&nbsp;&nbsp;{_inline(li)}', ps))
            elif nm == 'table':
                el = _make_table(c)
                if el:
                    target.append(el)
            elif nm in ('div', 'body', 'html', 'span', 'section', 'article'):
                st = _parse_style(c.get('style', ''))
                if st.get('page-break-before') == 'always':
                    target.append(_PB())
                if st.get('page-break-inside') == 'avoid':
                    sub = []
                    _walk(c, sub)
                    if sub:
                        target.append(_KT(sub))
                else:
                    _walk(c, target)
            else:
                _walk(c, target)

    soup = BeautifulSoup(html_content, 'html.parser')
    root = soup.body or soup
    _walk(root)
    if not flow:
        flow.append(Paragraph('&nbsp;', ParagraphStyle('e', fontName='Helvetica', fontSize=FONT_PT)))

    out = _io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4, leftMargin=MARGIN, rightMargin=MARGIN,
                            topMargin=MARGIN, bottomMargin=MARGIN)
    doc.build(flow)
    return out.getvalue()


@app.post("/api/export/pdf")
async def export_pdf(request: Request):
    """Convert HTML content to PDF and return as direct download (no print dialog)."""
    try:
        body = await request.json()
        html_content = body.get("html", "")
        filename = body.get("filename", "dokumen") or "dokumen"
        # Sanitize filename
        import re as _re
        filename = _re.sub(r'[^\w\-.]', '_', filename)
        if not filename.endswith(".pdf"):
            filename += ".pdf"
        if not html_content.strip():
            return JSONResponse({"ok": False, "error": "Tiada kandungan HTML."}, status_code=400)

        pdf_bytes = _render_html_to_pdf(html_content)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)


def _layout_html_to_docx_bytes(html_content: str) -> bytes:
    """Convert reconstruction HTML (paragraphs with text-align + <b>) to a .docx,
    preserving per-paragraph alignment and bold. Used to export an uploaded PDF
    as Word."""
    import io as _io2, re as _re2
    from bs4 import BeautifulSoup as _BS
    from docx import Document as _DocxDoc
    from docx.shared import Pt as _Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD

    soup = _BS(html_content, "html.parser")
    doc = _DocxDoc()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = _Pt(12)
    align_map = {"left": _WD.LEFT, "right": _WD.RIGHT, "center": _WD.CENTER, "justify": _WD.JUSTIFY}

    blocks = soup.find_all(["p", "h1", "h2", "h3", "h4", "li"])
    for el in blocks:
        text = el.get_text(separator=" ").strip()
        st = el.get("style") or ""
        if not text:
            doc.add_paragraph("")  # preserve blank-line spacing
            continue
        if el.name in ("h1", "h2", "h3", "h4"):
            para = doc.add_heading(text, level=int(el.name[1]))
        elif el.name == "li":
            para = doc.add_paragraph(text, style="List Bullet")
        else:
            para = doc.add_paragraph()
            m = _re2.search(r"text-align:\s*(left|right|center|justify)", st)
            if m:
                para.alignment = align_map.get(m.group(1))
            is_bold = bool(el.find(["b", "strong"])) or "font-weight:bold" in st.replace(" ", "")
            run = para.add_run(text)
            run.bold = is_bold
    buf = _io2.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_to_word_editable_bytes(raw_pdf) -> bytes:
    """Convert a PDF to an editable .docx that keeps the PDF's layout, using
    pdf2docx (reconstructs text, fonts, alignment and positioning via tables/
    text blocks — no page images, so the result stays editable)."""
    import io as _io, os as _os, tempfile as _tmp, logging as _log
    from pdf2docx import Converter
    _log.getLogger("pdf2docx").setLevel(_log.ERROR)  # silence verbose INFO output

    fd, tmp_path = _tmp.mkstemp(suffix=".docx")
    _os.close(fd)
    cv = Converter(stream=raw_pdf)
    try:
        cv.convert(tmp_path)
    finally:
        cv.close()
    with open(tmp_path, "rb") as f:
        data = f.read()
    try:
        _os.remove(tmp_path)
    except OSError:
        pass
    return data


@app.get("/api/review/download-word")
async def review_download_word(session_id: str = "default"):
    """Native-download endpoint: convert the session's uploaded PDF to an
    editable .docx and return it as an attachment. Triggered by a direct link
    click so the browser downloads it without a JS-gesture/blob workaround."""
    raw = _REVIEW_PDF_CACHE.get(session_id)
    if not raw:
        return JSONResponse({"error": "Tiada fail PDF dimuat naik untuk sesi ini."}, status_code=404)
    import io as _io
    try:
        try:
            docx_bytes = _pdf_to_word_editable_bytes(raw)
        except Exception:
            import pdfplumber
            with pdfplumber.open(_io.BytesIO(raw)) as pdf:
                html = _pdf_to_review_html(pdf)
            docx_bytes = _layout_html_to_docx_bytes(html) if html else None
        if not docx_bytes:
            return JSONResponse({"error": "Teks tidak dapat diekstrak dari PDF ini."}, status_code=400)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="dokumen.docx"'},
        )
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)


@app.post("/api/review/pdf-to-word")
async def review_pdf_to_word(file: UploadFile = File(...)):
    """Convert an uploaded PDF to an editable .docx that matches the PDF's layout.
    Uses pdf2docx; falls back to the paragraph reconstruction if that fails."""
    import io as _io
    try:
        raw = await file.read()
        try:
            docx_bytes = _pdf_to_word_editable_bytes(raw)
        except Exception:
            # Fallback: reconstruct from the layout HTML
            import pdfplumber
            with pdfplumber.open(_io.BytesIO(raw)) as pdf:
                html = _pdf_to_review_html(pdf)
            if not html:
                return JSONResponse({"error": "Teks tidak dapat diekstrak dari PDF ini."}, status_code=400)
            docx_bytes = _layout_html_to_docx_bytes(html)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="dokumen.docx"'},
        )
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)


@app.post("/api/review/download-edited")
async def review_download_edited(request: Request):
    """Convert edited HTML doc back to DOCX and return as download."""
    try:
        body = await request.json()
        html_content = body.get("html", "")
        if not html_content:
            return JSONResponse({"error": "No HTML content"}, status_code=400)

        import io as _io2
        from docx import Document as _DocxDoc2
        from docx.shared import Pt as _Pt

        # Try html2docx if available (best quality)
        try:
            from html2docx import html2docx as _html2docx
            buf = _html2docx(html_content, title="Dokumen Disemak")
            return Response(
                content=buf.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": 'attachment; filename="dokumen_disemak.docx"'},
            )
        except ImportError:
            pass

        # Fallback: BeautifulSoup → python-docx (text only, basic formatting)
        from bs4 import BeautifulSoup as _BS
        soup = _BS(html_content, "html.parser")

        doc = _DocxDoc2()
        # Set default font
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = _Pt(11)

        def _add_run_text(para, text, bold=False):
            run = para.add_run(text)
            if bold:
                run.bold = True

        for el in soup.find_all(["p", "h1", "h2", "h3", "h4", "li", "td"]):
            text = el.get_text(separator=" ").strip()
            if not text:
                continue
            if el.name in ("h1", "h2"):
                para = doc.add_heading(text, level=int(el.name[1]))
            elif el.name in ("h3", "h4"):
                para = doc.add_heading(text, level=int(el.name[1]))
            elif el.name == "li":
                para = doc.add_paragraph(text, style="List Bullet")
            else:
                para = doc.add_paragraph()
                bold_parts = el.find_all(["strong", "b"])
                if bold_parts:
                    raw = str(el)
                    # Simple bold detection: add whole text with bold if most is bold
                    _add_run_text(para, text, bold=len(bold_parts) > 0 and len(text) < 60)
                else:
                    para.add_run(text)

        buf = _io2.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="dokumen_disemak.docx"'},
        )
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)


@app.get("/api/data-status")
async def data_status(session_id: str = "default"):
    data = da_get_data(session_id)
    if data:
        return JSONResponse({
            "has_data": True,
            "filename": data["filename"],
            "shape": data["shape"],
            "columns": data["columns"],
        })
    return JSONResponse({"has_data": False})


@app.get("/api/document/download")
async def download_document(session_id: str = "default"):
    docx_bytes = lg_build_docx(session_id)
    if not docx_bytes:
        return JSONResponse({"error": "Tiada dokumen sedia untuk dimuat turun."}, status_code=404)
    info = lg_get_session(session_id)
    doc_type = info.get("doc_type", "dokumen") if info else "dokumen"
    filename = f"{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/report/download")
async def download_report(session_id: str = "default"):
    docx_bytes = rg_build_docx(session_id)
    if not docx_bytes:
        return JSONResponse({"error": "Tiada laporan sedia untuk dimuat turun."}, status_code=404)
    filename = f"laporan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/report/save")
async def save_report(req: "SaveDocRequest"):
    session = rg_get_session(req.session_id)
    if session:
        session["document"] = req.content
        return JSONResponse({"ok": True})
    return JSONResponse({"ok": False, "error": "Sesi tidak ditemui."}, status_code=404)


@app.post("/api/report/send-email")
async def send_report_email(req: "EmailRequest"):
    result = rg_send_email(req.session_id, req.to_email, req.subject)
    return JSONResponse(result)


@app.post("/api/report/upload-image")
async def report_upload_image(file: UploadFile = File(...), session_id: str = Form("default")):
    contents = await file.read()
    result = rg_add_image(session_id, contents, file.filename or "image.jpg")
    return JSONResponse(result)


@app.get("/api/report/images")
async def report_get_images(session_id: str = "default"):
    images = rg_get_images(session_id)
    return JSONResponse({
        "images": [
            {"index": i, "filename": img["filename"], "url": f"/api/report/image-file/{img['safe_name']}"}
            for i, img in enumerate(images)
        ],
        "count": len(images),
        "max": 4,
    })


@app.delete("/api/report/image/{index}")
async def report_delete_image(index: int, session_id: str = "default"):
    result = rg_remove_image(session_id, index)
    return JSONResponse(result)


@app.get("/api/report/image-file/{safe_name}")
async def report_image_file(safe_name: str):
    p = _Path("static/report_images") / safe_name
    if not p.exists():
        return JSONResponse({"error": "Fail tidak ditemui."}, status_code=404)
    return _FileResponse(str(p))


class SaveDocRequest(BaseModel):
    session_id: str = "default"
    content: str


@app.post("/api/document/save")
async def save_document(req: SaveDocRequest):
    from agents.letter_generator import get_session_info
    session = get_session_info(req.session_id)
    if session:
        session["document"] = req.content
        return JSONResponse({"ok": True})
    return JSONResponse({"ok": False, "error": "Sesi tidak ditemui."}, status_code=404)


class EmailRequest(BaseModel):
    session_id: str = "default"
    to_email: str
    subject: str


@app.post("/api/document/send-email")
async def send_document_email(req: EmailRequest):
    result = lg_send_email(req.session_id, req.to_email, req.subject)
    return JSONResponse(result)


class ExportRequest(BaseModel):
    session_id: str = "default"
    format: str = "pptx"
    data: dict
    chart_image: str | None = None


@app.post("/api/analysis/export")
async def export_analysis(req: ExportRequest):
    from agents.data_analysis_export import build_pptx, build_pdf
    try:
        if req.format == "pptx":
            content = build_pptx(req.data, chart_image_b64=req.chart_image)
            return Response(
                content=content,
                media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                headers={"Content-Disposition": 'attachment; filename="analisis_data.pptx"'},
            )
        elif req.format == "pdf":
            content = build_pdf(req.data, chart_image_b64=req.chart_image)
            return Response(
                content=content,
                media_type="application/pdf",
                headers={"Content-Disposition": 'attachment; filename="analisis_data.pdf"'},
            )
        else:
            return JSONResponse({"error": "Format tidak disokong."}, status_code=400)
    except Exception as e:
        return JSONResponse({"error": f"Gagal menjana fail: {e}"}, status_code=500)


from backend.letterhead_store import (
    list_letterheads,
    get_active_by_type,
    add_letterhead,
    set_active as lh_set_active,
    delete_letterhead,
    update_name as lh_update_name,
)
from fastapi.responses import FileResponse as _FileResponse
from pathlib import Path as _Path


@app.get("/api/letterhead/list")
async def letterhead_list():
    data = list_letterheads()
    active_lh = get_active_by_type("letterhead")
    active_logo = get_active_by_type("logo")
    return JSONResponse({
        "letterheads": data,
        "active_letterhead_id": active_lh["id"] if active_lh else None,
        "active_logo_id": active_logo["id"] if active_logo else None,
    })


@app.get("/api/letterhead/active")
async def letterhead_active(lh_type: str = "letterhead"):
    lh = get_active_by_type(lh_type)
    return JSONResponse({"active": lh})


@app.post("/api/letterhead/upload")
async def letterhead_upload(
    file: UploadFile = File(...),
    label: str = Form(""),
    lh_type: str = Form("letterhead"),
):
    ext = _Path(file.filename).suffix.lower()
    allowed = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}
    if ext not in allowed:
        return JSONResponse({"ok": False, "error": f"Format tidak disokong: {ext}"}, status_code=400)
    data = await file.read()
    try:
        entry = add_letterhead(data, file.filename, label, lh_type)
    except ValueError as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=400)
    return JSONResponse({"ok": True, "letterhead": entry})


class LhSelectRequest(BaseModel):
    id: str
    lh_type: str = "letterhead"


@app.post("/api/letterhead/select")
async def letterhead_select(req: LhSelectRequest):
    ok = lh_set_active(req.id, req.lh_type)
    return JSONResponse({"ok": ok})


@app.post("/api/letterhead/rename")
async def letterhead_rename(req: dict):
    ok = lh_update_name(req.get("id", ""), req.get("name", ""))
    return JSONResponse({"ok": ok})


@app.delete("/api/letterhead/{lh_id}")
async def letterhead_delete(lh_id: str):
    ok = delete_letterhead(lh_id)
    return JSONResponse({"ok": ok})


@app.get("/api/letterhead/image/{filename}")
async def letterhead_image(filename: str):
    p = _Path("static/letterheads") / filename
    if not p.exists():
        return JSONResponse({"error": "Fail tidak ditemui."}, status_code=404)
    return _FileResponse(str(p))


@app.post("/api/clear")
async def clear(req: ChatRequest):
    _get_store().clear_session(req.session_id)
    from agents.data_analysis import clear_session_data
    from agents.letter_generator import clear_session as lg_clear
    from agents.report_generator import clear_session as rg_clear
    from agents.kpm_support import clear_session as kpm_clear
    from agents.document_reviewer import clear_session as dr_clear
    clear_session_data(req.session_id)
    lg_clear(req.session_id)
    rg_clear(req.session_id)
    kpm_clear(req.session_id)
    dr_clear(req.session_id)
    return JSONResponse({"status": "cleared"})


@app.get("/api/agents")
async def list_agents():
    return JSONResponse(AGENT_LABELS)


@app.get("/api/health")
async def health_check():
    """Semak kesihatan sistem dan status circuit breaker DeepSeek."""
    from backend.deepseek_client import get_circuit_breaker_status
    cb = get_circuit_breaker_status()
    ok = cb["state"] != "OPEN"
    return JSONResponse({
        "status": "ok" if ok else "degraded",
        "llm": {
            "provider": "deepseek",
            "circuit_breaker": cb,
        },
    }, status_code=200 if ok else 503)


@app.post("/api/feedback")
async def submit_feedback(req: FeedbackRequest):
    _feedback.append({
        "session_id": req.session_id,
        "message_index": req.message_index,
        "feedback": req.feedback,
        "agent": req.agent,
        "timestamp": datetime.now().isoformat(),
    })
    return JSONResponse({"ok": True})


@app.get("/api/stats")
async def get_stats():
    store = _get_store()
    all_sessions = store.list_sessions()
    agent_counts: dict[str, int] = {}
    agent_messages: dict[str, int] = {}
    total_messages = 0
    for meta in all_sessions:
        agent = meta.get("agent", "unknown")
        agent_counts[agent] = agent_counts.get(agent, 0) + 1
        mc = meta.get("msg_count", 0)
        agent_messages[agent] = agent_messages.get(agent, 0) + mc
        total_messages += mc

    fb_by_agent: dict[str, dict] = {}
    for f in _feedback:
        a = f.get("agent", "unknown")
        if a not in fb_by_agent:
            fb_by_agent[a] = {"up": 0, "down": 0}
        fb_by_agent[a][f["feedback"]] = fb_by_agent[a].get(f["feedback"], 0) + 1

    recent = all_sessions[:10]
    return JSONResponse({
        "total_sessions": len(all_sessions),
        "total_messages": total_messages,
        "agent_counts": agent_counts,
        "agent_messages": agent_messages,
        "feedback_total": {"up": sum(f["feedback"] == "up" for f in _feedback), "down": sum(f["feedback"] == "down" for f in _feedback)},
        "feedback_by_agent": fb_by_agent,
        "recent_sessions": [{"session_id": s["session_id"], **s} for s in recent],
        "agent_labels": {k: v["name"] for k, v in AGENT_LABELS.items()},
    })


@app.get("/api/sessions")
async def list_sessions(agent: str = ""):
    sessions = _get_store().list_sessions(agent=agent)
    return JSONResponse(sessions)


@app.get("/api/sessions/{session_id}/messages")
async def get_session_messages(session_id: str):
    history = _get_store().get_messages(session_id)
    messages = []
    for msg in history:
        agent = msg.get("agent", "")
        agent_info = AGENT_LABELS.get(agent, {"icon": "", "name": ""})
        messages.append({
            "role": msg["role"],
            "content": msg["content"],
            "agent": agent,
            "agent_icon": agent_info.get("icon", ""),
            "agent_name": agent_info.get("name", ""),
        })
    return JSONResponse(messages)


@app.post("/api/agent-chat")
async def agent_chat(req: AgentChatRequest, request: Request):
    """Direct chat with a specific agent, bypassing intent classification."""
    user = get_current_user(request)
    user_name = ""
    if user:
        profile = get_profile(user["sub"])
        user_name = profile.get("nama") or user.get("name", "")
    store = _get_store()
    is_intro = req.message == "__INTRO__"
    history = store.get_messages(req.session_id)
    if not is_intro:
        store.append_message(req.session_id, {"role": "user", "content": req.message})
        history.append({"role": "user", "content": req.message})

    agent = req.agent
    agent_info = AGENT_LABELS.get(agent, {"icon": "\U0001f4ac", "name": agent})

    try:
        from agents import data_analysis, letter_generator, kpm_support, report_generator, document_reviewer
        import inspect
        agent_map = {
            "data_analysis": data_analysis,
            "letter_generator": letter_generator,
            "kpm_support": kpm_support,
            "report_generator": report_generator,
            "document_reviewer": document_reviewer,
        }
        module = agent_map.get(agent)
        if module and hasattr(module, "handle"):
            sig = inspect.signature(module.handle)
            kwargs = {"query": req.message, "history": history}
            if "session_id" in sig.parameters:
                kwargs["session_id"] = req.session_id
            if "lang" in sig.parameters:
                kwargs["lang"] = req.lang
            if "user_name" in sig.parameters:
                kwargs["user_name"] = user_name
            output = module.handle(**kwargs)

            # JSON retry: if agent expects JSON but output fails to parse, try once more
            JSON_AGENTS = {"data_analysis", "letter_generator", "report_generator", "document_reviewer"}
            if agent in JSON_AGENTS and not is_intro:
                try:
                    json.loads(output)
                except (json.JSONDecodeError, ValueError):
                    sys_note = "[SYSTEM: Reply ONLY in valid JSON format. No text outside JSON.]" if req.lang == "en" else "[SISTEM: Balas HANYA dalam format JSON yang sah. Tiada teks lain di luar JSON.]"
                    retry_kwargs = {**kwargs, "query": req.message + f"\n\n{sys_note}"}
                    try:
                        output = module.handle(**retry_kwargs)
                    except Exception:
                        pass
        else:
            output = f"Agen '{agent}' tidak ditemui."
    except Exception as e:
        output = f"Ralat: {e}"

    if not is_intro:
        store.append_message(req.session_id, {"role": "assistant", "content": output, "agent": agent})
        existing_meta = store.get_meta(req.session_id) or {}
        store.upsert_meta(
            req.session_id,
            agent=agent,
            title=existing_meta.get("title") or req.message[:60],
        )

    structured = None
    if agent in ("data_analysis", "letter_generator", "report_generator", "document_reviewer"):
        structured = _parse_agent_json(output)

    # For an uploaded PDF under review, locate each issue's text and attach a
    # highlight box so the frontend can mark the errors on the page images.
    if agent == "document_reviewer" and structured and structured.get("issues"):
        raw_pdf = _REVIEW_PDF_CACHE.get(req.session_id)
        if raw_pdf:
            try:
                _locate_issue_boxes(raw_pdf, structured["issues"])
                output = json.dumps(structured, ensure_ascii=False)
            except Exception:
                pass

    # Auto-review + auto-improve: when a document agent produces a ready document
    _DOC_AGENTS = {
        "letter_generator": lambda s: "Memo Dalaman" if s.get("doc_type") == "memo" else "Surat Rasmi",
        "report_generator": lambda s: "Laporan Satu Muka Surat",
    }
    if agent in _DOC_AGENTS and structured and structured.get("ready_to_save") and structured.get("document_preview"):
        try:
            from agents.document_reviewer import auto_review, auto_improve
            doc_text = structured["document_preview"]
            doc_type = _DOC_AGENTS[agent](structured)

            # Step 1: Review
            review = auto_review(doc_text, doc_type, req.session_id, req.lang)
            if review:
                structured["auto_review"] = review

                # Step 2: Auto-improve using review findings
                agent_mod = __import__(f"agents.{agent}", fromlist=[agent])
                all_fields = agent_mod.get_fields(req.session_id)
                gen_keys = agent_mod.GENERATED_FIELD_KEYS

                # Strip template-hardcoded opening from surat isi before sending to LLM
                _SURAT_ISI_PREFIX = "dengan hormatnya perkara di atas adalah dirujuk"
                clean_fields = dict(all_fields)
                if agent == "letter_generator" and "isi" in clean_fields:
                    isi_lines = clean_fields["isi"].strip().splitlines()
                    if isi_lines and isi_lines[0].strip().lower().rstrip(".") == _SURAT_ISI_PREFIX:
                        clean_fields["isi"] = "\n".join(isi_lines[1:]).strip()

                improvement = auto_improve(
                    doc_type=doc_type,
                    user_fields=clean_fields,
                    generated_field_keys=gen_keys,
                    review=review,
                    lang=req.lang,
                )
                if improvement and improvement.get("improved_fields"):
                    # Strip template prefix from improved isi if LLM included it
                    imp_fields = improvement["improved_fields"]
                    if agent == "letter_generator" and "isi" in imp_fields:
                        isi_lines = imp_fields["isi"].strip().splitlines()
                        if isi_lines and isi_lines[0].strip().lower().rstrip(".") == _SURAT_ISI_PREFIX:
                            imp_fields["isi"] = "\n".join(isi_lines[1:]).strip()
                    # apply_improvement updates fields + rebuilds doc via original template
                    new_doc = agent_mod.apply_improvement(req.session_id, imp_fields)
                    if new_doc:
                        structured["document_preview"] = new_doc
                        # Rebuild memo HTML from updated fields
                        if agent == "letter_generator" and structured.get("doc_type") == "memo":
                            from agents.letter_generator import _build_memo_html
                            new_fields = agent_mod.get_fields(req.session_id)
                            structured["document_html"] = _build_memo_html(new_fields)
                    structured["auto_review"]["improvement"] = {
                        "changes_applied": improvement.get("changes_applied", []),
                        "changes_skipped": improvement.get("changes_skipped", []),
                        "needs_info": improvement.get("needs_info"),
                    }

                output = json.dumps(structured, ensure_ascii=False)
        except Exception:
            pass

    _shorten_doc_message(structured, getattr(req, "lang", "ms"))

    return JSONResponse({
        "response": output,
        "agent": agent,
        "agent_icon": agent_info["icon"],
        "agent_name": agent_info["name"],
        "structured": structured,
    })


@app.get("/api/avatar")
async def api_avatar(request: Request):
    """Proxy Google profile picture server-side to avoid browser CORS/referrer blocks."""
    user = get_current_user(request)
    if not user or not user.get("picture"):
        return Response(status_code=204)
    pic_url = user["picture"]
    # Request larger size for retina displays
    pic_url = pic_url.split("=s")[0] + "=s120-c" if "=" in pic_url else pic_url
    try:
        import httpx
        async with httpx.AsyncClient(follow_redirects=True, timeout=8.0) as client:
            resp = await client.get(pic_url, headers={"User-Agent": "Mozilla/5.0"})
        if resp.status_code == 200:
            content_type = resp.headers.get("content-type", "image/jpeg")
            return Response(content=resp.content, media_type=content_type,
                            headers={"Cache-Control": "public, max-age=3600"})
    except Exception:
        pass
    return Response(status_code=204)


@app.get("/api/profile")
async def api_get_profile(request: Request):
    user = get_current_user(request)
    if not user:
        return JSONResponse({"error": "Tidak log masuk."}, status_code=401)
    profile = get_profile(user["sub"])
    return JSONResponse({
        "nama":    profile.get("nama", "") or user.get("name", ""),
        "jawatan": profile.get("jawatan", ""),
        "stesen":  profile.get("stesen", ""),
        "daerah":  profile.get("daerah", ""),
        "negeri":  profile.get("negeri", ""),
        "email":   user.get("email", ""),
        "picture": user.get("picture", ""),
    })


@app.post("/api/profile")
async def api_save_profile(request: Request):
    user = get_current_user(request)
    if not user:
        return JSONResponse({"error": "Tidak log masuk."}, status_code=401)
    data = await request.json()
    profile = save_profile(user["sub"], user["email"], data)
    return JSONResponse({"ok": True, "profile": dict(profile)})


if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8112, reload=True)
